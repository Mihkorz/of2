# -*- coding: utf-8 -*-
import os
from datetime import datetime
import csv
import json
import cairosvg
import numpy as np
from scipy.stats.mstats import gmean
from scipy.stats import ttest_1samp
from pandas import read_csv, read_excel, DataFrame, Series
from docx import Document as pyDocx
from docx.shared import Inches
from cStringIO import StringIO

from django.http import HttpResponse
from django.views.generic.detail import DetailView
from django.views.generic.base import TemplateView
from django.views.generic.edit import FormView 
from django.views.generic import ListView
from django.http import HttpResponseRedirect
from django.core.urlresolvers import reverse
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from django.conf import settings
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile

from .models import Nosology, TreatmentMethod
from database.models import Pathway 
from profiles.models import Document
from core.forms import MedicCalculationParametersForm
from core.stats import quantile_normalization, XPN_normalisation, fdr_corr


class MedicNosologyList(ListView):
    """
    List of Nosologies for Medic DB section
    """    
    model = Nosology
    template_name = 'medic/medic_nosology_list.html'
    context_object_name = 'nosologies'
    paginate_by = 20
        
    def get_context_data(self, **kwargs):
        context = super(MedicNosologyList, self).get_context_data(**kwargs)        
        return context
        
    @method_decorator(login_required)
    def dispatch(self, request, *args, **kwargs):
        return super(MedicNosologyList, self).dispatch(request, *args, **kwargs)
    
class MedicNosologyDetail(DetailView):
    """
    Details page for particular nosology
    """
    model = Nosology
    template_name = 'medic/medic_nosology_detail.html'
    context_object_name = 'nosology'    
    
    @method_decorator(login_required)
    def dispatch(self, request, *args, **kwargs):
        return super(MedicNosologyDetail, self).dispatch(request, *args, **kwargs)
    
    def get_context_data(self, **kwargs):
        context = super(MedicNosologyDetail, self).get_context_data(**kwargs)
        
        return context


class Sample(object):
    responce_score = 0
    initial_status = ''
    status_by_voting = ''
    status_by_window = ''
    
        
class MedicTreatmentDetail(DetailView):
    """
    Details page for particular treatment
    """
    model = TreatmentMethod
    template_name = 'medic/medic_treatment_detail.html'
    context_object_name = 'treatment'    
    
    @method_decorator(login_required)
    def dispatch(self, request, *args, **kwargs):
        return super(MedicTreatmentDetail, self).dispatch(request, *args, **kwargs)
    
    def get_context_data(self, **kwargs):
        context = super(MedicTreatmentDetail, self).get_context_data(**kwargs)
        
        file_pms1 = settings.MEDIA_ROOT+"/"+self.object.file_pms1.name
        file_probability = settings.MEDIA_ROOT+"/"+self.object.file_probability.name
        
        
        sniffer = csv.Sniffer()
        dialect = sniffer.sniff(open(file_probability, 'r').read(), delimiters='\t,;') # defining the separator of the csv file
        df_prob = read_csv(file_probability, delimiter=dialect.delimiter)
        grouped = df_prob.groupby('Sample', sort=True)
                
        all_samples = len(df_prob.index)/2
                
        lResponders = []
        lnonResponders = []
        lAllRespScore = []
        lSampleObjects = []
        
        true_positive = 0
        true_negative = 0
        false_positive = 0
        false_negative = 0
        
        for name, group in grouped:
            status = "NRES" if "NRES" in name else "RES"
            path_cols = [col for col in group.columns if col not in ['Sample', 'group']]
            path_df = group[path_cols]
            
            divided_df = path_df.T[path_df.index[1]] / path_df.T[path_df.index[0]]
            Rcount = 0 # count Responder votes
            for index, val in divided_df.iteritems():
                if val>1:
                    Rcount+= 1
            
            ratio = float(Rcount)/float(len(path_cols))
            if ratio > 0.5 and status == 'RES':
                true_positive+=1
            if ratio > 0.5 and status == 'NRES':
                false_negative+=1
            if ratio <= 0.5 and status == 'NRES':
                true_negative+=1
            if ratio <= 0.5 and status == 'RES':
                false_positive+=1                
            
            lAllRespScore.append(ratio)
            
            sampleObj = Sample() # Create sample object to test 0.1 window method
            sampleObj.responce_score = ratio
            sampleObj.initial_status = status
            lSampleObjects.append(sampleObj)
            
            if status == 'NRES': 
                lnonResponders.append(ratio)
            else:
                lResponders.append(ratio)
        
        from collections import Counter, OrderedDict
        
        nresponders = dict(Counter(lnonResponders))
        for x in nresponders:
            nresponders[x]/=float(all_samples)
            nresponders[x]*=100
        responders =  dict(Counter(lResponders))
        for x in responders:
            responders[x]/=float(all_samples)
            responders[x]*=100
        
        dAccur = {}
        dSp = {}
        dSn= {}
        dBalnc={} 
        for score in lAllRespScore:
            true_positive = 0
            true_negative = 0
            false_positive = 0
            false_negative = 0
            for name, group in grouped:
                status = "NRES" if "NRES" in name else "RES"
                path_cols = [col for col in group.columns if col not in ['Sample', 'group']]
                path_df = group[path_cols]
            
                divided_df = path_df.T[path_df.index[1]] / path_df.T[path_df.index[0]]
                Rcount = 0 # count Responder votes
                for index, val in divided_df.iteritems():
                    if val>1:
                        Rcount+= 1
                ratio = float(Rcount)/float(len(path_cols))
                if ratio >= score and status == 'RES':
                    true_positive+=1
                if ratio > score and status == 'NRES':
                    false_negative+=1
                if ratio <= score and status == 'NRES':
                    true_negative+=1
                if ratio <= score and status == 'RES':
                    false_positive+=1
                try:    
                    specificity = (float(true_negative))/(false_positive+true_negative)
                except: 
                    specificity = 0.0001
                try:
                    sensitivity = (float(true_positive))/(true_positive+false_negative)
                except:
                    sensitivity = 0.0001
                try:
                    AUC = (specificity+sensitivity)/2
                except:
                    AUC = 0
                try:
                    accuracy = (true_positive+true_negative)/float(all_samples)
                except:
                    accuracy = 0
                dSp[score] = specificity
                dSn[score] = sensitivity
                dBalnc[score] = AUC
                dAccur[score] = accuracy
            #raise Exception('true')
            
        
        """ Statistical values """
        try:
            specificity = (float(true_negative))/(false_positive+true_negative)
        except:
            specificity = 0.00001
        try:
            sensitivity = (float(true_positive))/(true_positive+false_negative)
        except: sensitivity = 0.00001
        AUC = (specificity+sensitivity)/2
        accuracy = (true_positive+true_negative)/float(all_samples)
               
        context['nres'] = OrderedDict(sorted(nresponders.items()))
        context['res'] = OrderedDict(sorted(responders.items()))                
        
        context['specificity'] = specificity
        context['sensitivity'] = sensitivity
        context['AUC'] = AUC
        context['accuracy'] = accuracy
        
        """ getting best statistaical values for different score"""
        import operator
        bestSP = max(dSp.iteritems(), key=operator.itemgetter(1))[0]
        bestSn = max(dSn.iteritems(), key=operator.itemgetter(1))[0]
        bestBalanced = max(dBalnc.iteritems(), key=operator.itemgetter(1))[0]
        bestAccur = max(dAccur.iteritems(), key=operator.itemgetter(1))[0]
        
        context['dSn'] = OrderedDict(sorted(dSn.items()))
        context['dSp'] = OrderedDict(sorted(dSp.items()))
        context['dBalnc'] = OrderedDict(sorted(dBalnc.items()))
        context['dAccur'] = OrderedDict(sorted(dAccur.items()))
        
        context['bSPkey'] = bestSP
        context['bSPVal'] = dSp[bestSP]
        context['bSnkey'] = bestSn
        context['bSnVal'] = dSn[bestSn]
        context['bBalkey'] = bestBalanced
        context['bBalVal'] = dBalnc[bestBalanced]
        context['bAcckey'] = bestAccur
        context['bAccVal'] = dAccur[bestAccur]
        
        
        """ 0.1 window method """
        lnewResp = []
        lnewNonResp = []
        ttt = []
        
        for rangee in np.arange(0.1,1.1,0.1):
            ttt.append(rangee)
            count_responders = 0
            count_nresponders = 0
            filtered = [x for x in lSampleObjects if (x.responce_score>=rangee-0.1 and x.responce_score<=rangee)]
            for s in filtered:
                if s.initial_status == "RES":
                    count_responders+=1
                else:
                    count_nresponders+=1
            if count_nresponders>=count_responders:
                for x in filtered:
                    x.status_by_window = "NRES"
                    lnewNonResp.append(x.responce_score) 
            else:
                for x in filtered:
                    x.status_by_window = "RES"
                    lnewResp.append(x.responce_score)
                
            
        new_nresponders = dict(Counter(lnewNonResp))
        for x in new_nresponders:
            new_nresponders[x]/=float(all_samples)
            new_nresponders[x]*=100
        new_responders =  dict(Counter(lnewResp))
        for x in new_responders:
            new_responders[x]/=float(all_samples)
            new_responders[x]*=100
        
        true_positive = 0
        true_negative = 0
        false_positive = 0
        false_negative = 0    
        for sample in lSampleObjects:
            if sample.initial_status =='RES' and sample.status_by_window == 'RES':
                    true_positive+=1
            if sample.initial_status =='RES' and sample.status_by_window == 'NRES':
                    false_negative+=1
            if sample.initial_status =='NRES' and sample.status_by_window == 'NRES':
                    true_negative+=1
            if sample.initial_status =='NRES' and sample.status_by_window == 'RES':
                    false_positive+=1
                    
        try:
            new_specificity = (float(true_negative))/(false_positive+true_negative)
        except: 
            new_specificity = 0.0001
        try:
            new_sensitivity = (float(true_positive))/(true_positive+false_negative)
        except:
            new_sensitivity = 0.0001
        try:
            new_AUC = (new_specificity+new_sensitivity)/2
        except:
            new_AUC = 0
        try:
            new_accuracy = (true_positive+true_negative)/float(all_samples)
        except:
            new_accuracy = 0 
            
        context['new_nres'] = OrderedDict(sorted(new_nresponders.items()))
        context['new_res'] = OrderedDict(sorted(new_responders.items()))      
        
        context['new_specificity'] = new_specificity
        context['new_sensitivity'] = new_sensitivity
        context['new_AUC'] = new_AUC
        context['new_accuracy'] = new_accuracy
                
        
        df_pms1 = read_csv(file_pms1)
        context['PAS1'] = df_pms1.to_html()
        context['prob'] = df_prob.to_html()
        #raise Exception("exp")
        return context
    
class PatientTreatmentDetail(DetailView):
    """ 
    Details for patient for each treatment
    """
    model = Document
    template_name = 'medic/patient_treatment_detail.html'
    
    @method_decorator(login_required)
    def dispatch(self, request, *args, **kwargs):
        return super(PatientTreatmentDetail, self).dispatch(request, *args, **kwargs)
    
    def get_context_data(self, **kwargs):
        context = super(PatientTreatmentDetail, self).get_context_data(**kwargs)
        
        treatment = TreatmentMethod.objects.get(pk=self.kwargs['treat_id'])
        
        filename = settings.MEDIA_ROOT+"/"+self.object.document.name
        filename = filename.replace('patient_results', treatment.name)        
        
        
        file_pms1 = settings.MEDIA_ROOT+"/"+treatment.file_pms1.name
        file_probability = settings.MEDIA_ROOT+"/"+treatment.file_probability.name
        
        sniffer = csv.Sniffer()
        dialect = sniffer.sniff(open(file_probability, 'r').read(), delimiters='\t,;') # defining the separator of the csv file
        df_prob = read_csv(file_probability, delimiter=dialect.delimiter)
        dialect = sniffer.sniff(open(filename, 'r').read(), delimiters='\t,;')
        df_output = read_csv(filename, delimiter=dialect.delimiter)        
        
        path_cols = [col for col in df_prob.columns if col not in ['Sample', 'group', 'Unnamed: 0']]
        path_num = len(path_cols)
                
        df_pms1 = read_csv(file_pms1, index_col="Pathway").transpose()
        df_required_paths = df_pms1[path_cols]
        df_required_paths.reset_index(inplace="True")
        df_nres = df_required_paths[df_required_paths['index'].str.contains("NRES")]
        df_res = df_required_paths[~df_required_paths['index'].str.contains("NRES")]
        
        patient_responder = {}
        patient_nonresponder = {}
        
        for path_name in path_cols:
            from scipy.stats import norm 
            if path_name in df_output.index:
                patient_pms1 = df_output.loc[path_name.strip()].item()
            else:
                patient_pms1 = 0
                
            r_mean = df_res[path_name].mean(axis=1)
            r_std = df_res[path_name].std(axis=1)
            
            if patient_pms1 <= r_mean:
                    r_probability = norm.cdf(patient_pms1, r_mean, r_std)
            else:
                    r_probability = 1- norm.cdf(patient_pms1, r_mean, r_std)
                    
            patient_responder[path_name] = r_probability
            
            nr_mean = df_nres[path_name].mean(axis=1)
            nr_std = df_nres[path_name].std(axis=1)
            
            if patient_pms1 <= nr_mean:
                    nr_probability = norm.cdf(patient_pms1, nr_mean, nr_std)
            else:
                    nr_probability = 1- norm.cdf(patient_pms1, nr_mean, nr_std)
                    
            patient_nonresponder[path_name] = nr_probability
            
        dict_for_df={}
        dict_for_df['Responder'] = patient_responder
        dict_for_df['Nonresponder'] = patient_nonresponder
        
        df_patient = DataFrame.from_dict(dict_for_df).transpose()
        
        divided_df_patient = df_patient.T[df_patient.index[1]] / df_patient.T[df_patient.index[0]]
        patient_Rcount = 0
        for index, val in divided_df_patient.iteritems():
            if val>1:
                patient_Rcount+=1
        ratio = float(patient_Rcount)/float(path_num)
        
        lResponders = []
        lnonResponders = []
        flag_responder = 1
        patient_votes = float(patient_Rcount)/float(path_num)
        if ratio > 0.5:
            lResponders.append(patient_votes )
        else:
            lnonResponders.append(patient_votes)
            flag_responder = 0       
                
        
        """ DRAW HISTOGRAM"""
        grouped = df_prob.groupby('Sample', sort=True)
        
        df_c = df_prob[df_prob['Sample'].str.contains("NRES")]
        all_samples = len(df_prob.index)/2
        num_nres_samples = len(df_c.index)/2 + len(lnonResponders)
        num_res_samples = (all_samples - num_nres_samples) +len(lResponders)
        
        num_responders = 0 # number of responders
        num_guessed_rigth = 0 # number of responders that remained status after voting
        num_guessed_wrong = 0 # number of samples that changed status after voting
        
        true_positive = 0
        true_negative = 0
        false_positive = 0
        false_negative = 0
        
        for name, group in grouped:
            status = "NRES" if "NRES" in name else "RES"
            path_cols = [col for col in group.columns if col not in ['Sample', 'group', 'Unnamed: 0']]
            path_df = group[path_cols]
            #raise
            divided_df = path_df.T[path_df.index[1]] / path_df.T[path_df.index[0]]
            Rcount = 0 # count Responder votes
            for index, val in divided_df.iteritems():
                if val>1:
                    Rcount+= 1
            ratio = float(Rcount)/float(len(path_cols))
            if status == 'RES':
                num_responders+=1
            
            if ratio > 0.5 and status == 'RES':
                num_guessed_rigth+=1
                true_positive+=1
            if ratio > 0.5 and status == 'NRES':
                num_guessed_wrong+=1
                false_negative+=1
            if ratio <= 0.5 and status == 'NRES':
                num_guessed_rigth+=1
                true_negative+=1
            if ratio <= 0.5 and status == 'RES':
                num_guessed_wrong+=1
                false_positive+=1 
            
            len_p = len(path_cols)
            if status == 'NRES': 
                lnonResponders.append(float(Rcount)/float(len_p))
            else:
                lResponders.append(float(Rcount)/float(len_p) )
        from collections import Counter, OrderedDict
        
        nresponders = dict(Counter(lnonResponders))
        for x in nresponders:
            nresponders[x]/=float(all_samples)
            nresponders[x]*=100
        responders =  dict(Counter(lResponders))
        for x in responders:
            responders[x]/=float(all_samples)
            responders[x]*=100
               
        
        """ Statistical values """
        try:
            specificity = (float(true_negative))/(false_positive+true_negative)
        except:
            specificity = 'Not difined.'
        try:    
            sensitivity = (float(true_positive))/(true_positive+false_negative)
        except:
            specificity = 'Not difined.'
        AUC = (specificity+sensitivity)/2
        accuracy = (true_positive+true_negative)/float(all_samples)
        
        context['nres'] = OrderedDict(sorted(nresponders.items()))
        context['res'] = OrderedDict(sorted(responders.items()))
        context['flag_responder'] = flag_responder
        context['patient_votes'] = patient_votes
        
        context['specificity'] = specificity
        context['sensitivity'] = sensitivity
        context['AUC'] = AUC
        context['accuracy'] = accuracy 
        
        
        treatment.treatment = treatment.treatment.replace('\n', ' ').replace('\r', '')
        
        
        context['treatment'] = treatment
        
        return context
    
    def render_to_json_response(self, context, **response_kwargs):
        data = json.dumps(context)
        response_kwargs['content_type'] = 'application/json'
        return HttpResponse(data, **response_kwargs)
    
    def post(self, request, *args, **kwargs):
        if self.request.is_ajax():
            self.object = self.get_object()
            context = self.get_context_data(object=self.object)
            lnres =  []
            for key, value in context['nres'].items():
                djson = {}
                djson.clear()
                djson['x']=key
                djson['y']=value
                if key == context['patient_votes'] and context['flag_responder']== 0:
                    djson['color']='red'
                lnres.append(djson)
            lres = []    
            for key, value in context['res'].items():
                djson = {}
                djson.clear()
                djson['x']=key
                djson['y']=value
                if key == context['patient_votes'] and context['flag_responder']== 1:
                    djson['color']='red'
                lres.append(djson)
                    
                    
            if context['flag_responder'] > 0:
                responder = 'responder'
            else:
                responder = 'non-responder'
                    
            data = {
                'nres': lnres,
                'res': lres,
                'treatment': context['treatment'].treatment,
                'patient_votes': round(context['patient_votes'], 2),
                'specificity': round(context['specificity'], 2),
                'sensitivity': round(context['sensitivity'], 2),
                'AUC': round(context['AUC'], 2),
                'accuracy': round(context['accuracy'], 2),
                'responder': responder,
                'num_of_patients': context['treatment'].num_of_patients,
                'histological_type': context['treatment'].histological_type,
                'grade': context['treatment'].grade,
                'hormone_receptor_status': context['treatment'].hormone_receptor_status,
                'her2_status': context['treatment'].her2_status,
                'stage': context['treatment'].stage,
                'citation': context['treatment'].citation,
                'organization_name': context['treatment'].organization_name,
                }
        
        return self.render_to_json_response(data)
    
class MedicPatientCalculation(FormView):
    form_class = MedicCalculationParametersForm
    
    @method_decorator(login_required)
    def dispatch(self, request, *args, **kwargs):
        return super(MedicPatientCalculation, self).dispatch(request, *args, **kwargs)
    
    def get_context_data(self, **kwargs):
        context = super(MedicPatientCalculation, self).get_context_data(**kwargs)
        
        
        return context
    
    def form_valid(self, form):
                
        
        input_document = Document.objects.get(pk=int(self.request.POST.get('doc_id')))
        hormone_status = form.cleaned_data.get('hormone_status', False)
        her2_status = form.cleaned_data.get('her2_status', False)
        
        treatments = TreatmentMethod.objects.filter(nosology=input_document.project.nosology)
        if hormone_status!='0':
                    treatments = treatments.filter(hormone_receptor_status=hormone_status)
        if her2_status!='0':
                    treatments = treatments.filter(her2_status=her2_status)
        
        process_doc_df = read_csv(settings.MEDIA_ROOT+"/"+input_document.input_doc.document.name,
                                  sep='\t', index_col='SYMBOL').fillna(0)
        original_columns = process_doc_df.columns
        
        for treatment in treatments:
            sniffer = csv.Sniffer()
            res_file = treatment.file_res
            dialect = sniffer.sniff(res_file.read(), delimiters='\t,;')
            res_file.seek(0)
            res_df = read_csv(res_file, delimiter=dialect.delimiter,
                          index_col='SYMBOL') #create DataFrame for responders
                
            nres_file = treatment.file_nres
            dialect = sniffer.sniff(nres_file.read(), delimiters='\t,;')
            nres_file.seek(0)
            nres_df = read_csv(nres_file, delimiter=dialect.delimiter,
                           index_col='SYMBOL') #create DataFrame for non-responderss
            
            res_nres_joined = res_df.join(nres_df, how='inner')
            
            """ Performing XPN between norms and responders+non-responders """
            try:
                df_after_xpn = XPN_normalisation(res_nres_joined, process_doc_df, iterations=10)
            except:
                raise
             
            #df_after_xpn = read_csv(settings.MEDIA_ROOT+"/xpn_done.csv", index_col='SYMBOL')
            
            """ calculating PAS1 for each treatment """
            
            df_for_pas1 = df_after_xpn[original_columns]
            norms_df = df_for_pas1[[norm for norm in [col for col in df_for_pas1.columns if 'Norm' in col]]]
            log_norms_df = np.log(norms_df)#use this for t-test, assuming log(norm) is distributed normally
            s_mean_norm = norms_df.apply(gmean, axis=1) #series of mean norms for CNR
        
            df_for_pas1.drop(norms_df.columns, axis=1, inplace=True)
            
            def apply_filter(col):
                _, p_value = ttest_1samp(log_norms_df, np.log(col), axis=1)
                s_p_value = Series(p_value, index = col.index).fillna(0)
                fdr_q_values = fdr_corr(np.array(s_p_value))
                col = col[(fdr_q_values<0.05)]
                return col
        
            df_for_pas1 = df_for_pas1.apply(apply_filter, axis=0).fillna(0)
        
            df_for_pas1 = df_for_pas1.divide(s_mean_norm, axis=0).fillna(0) #acquire CNR values
            df_for_pas1.replace(0, 1, inplace=True)
            df_for_pas1 = np.log(df_for_pas1) # now we have log(CNR)
        
            pas1_all_paths = DataFrame()
            
            for pathway in Pathway.objects.all().prefetch_related('gene_set'):
                genes = DataFrame(list(pathway.gene_set.all()
                                      .values('name', 'arr'))).set_index('name')#fetch genes 
            
                genes.index.name = 'SYMBOL'
                
                joined = genes.join(df_for_pas1, how='inner').drop(['arr'], axis=1) 
                pas1_for_pathway = joined.apply(lambda x: x*genes['arr'].astype('float')).sum()            
                pas1_for_pathway = pas1_for_pathway.set_value('Pathway', pathway.name)
                pas1_for_pathway = DataFrame(pas1_for_pathway).T.set_index('Pathway')
            
                pas1_all_paths = pas1_all_paths.append(pas1_for_pathway)
                
            """ save pas1 file as treatment name """
            path = os.path.join('users', str(input_document.project.owner),
                                            str(input_document.project),'output')
            file_pas1 = default_storage.save(path+"/"+treatment.name+".csv", ContentFile(''))
            pas1_all_paths.to_csv(settings.MEDIA_ROOT+"/"+file_pas1)                                           
                       
            
            
            
            
        output_doc = Document()
        output_doc.doc_type = 2
        json_db = 'Human'
        output_doc.parameters = {
                                 'use_fdr': True,
                                 'use_ttest_1sam': True,
                                 'norm_algirothm': 'geometric',
                                 'db': json_db,
                                 'hormone_status': hormone_status,
                                 'her2_status': her2_status }
        output_doc.project = input_document.project
        output_doc.created_by = self.request.user
        output_doc.created_at = datetime.now()     
        output_doc.document = path+"/patient_results.csv"
        output_doc.related_doc = input_document
        output_doc.save()
        return HttpResponseRedirect(reverse('document_detail', args=(output_doc.id,)))


class PatientTreatmentPDF(DetailView):
    """ 
    Details for all treatments as pdf
    """
    model = Document
    template_name = 'medic/patient_treatment_pdf.html'
    
    @method_decorator(login_required)
    def dispatch(self, request, *args, **kwargs):
        return super(PatientTreatmentPDF, self).dispatch(request, *args, **kwargs)
    
    def get_context_data(self, **kwargs):
        context = super(PatientTreatmentPDF, self).get_context_data(**kwargs)
        
        hormone_status = self.object.parameters['hormone_status']
        her2_status = self.object.parameters['her2_status']
        
        treatments = TreatmentMethod.objects.filter(nosology=self.object.project.nosology)
        
        if hormone_status!='0':
                    treatments = treatments.filter(hormone_receptor_status=hormone_status)
                    
        if her2_status!='0':
                    treatments = treatments.filter(her2_status=her2_status)
        
        context['treatments'] = treatments
        
        return context    
    
    
class MedicAjaxGenerateReport(TemplateView):
    template_name = 'medic/patient_treatment_pdf.html'
    
    def post(self, request, *args, **kwargs):
        
        t_name = request.POST.get('t_name')
        t_n_patients = request.POST.get('t_n_patients')
        t_hist = request.POST.get('t_hist')
        t_grade = request.POST.get('t_grade')
        t_hormone = request.POST.get('t_hormone')
        t_her2 = request.POST.get('t_her2')
        t_stage = request.POST.get('t_stage')
        t_treatment = request.POST.get('t_treatment')
        t_cit = request.POST.get('t_cit')
        o_org = request.POST.get('o_org')
        svg = request.POST.get('svg')
        
        
        
        fout = open(settings.MEDIA_ROOT+'/output.png','w')
        
        cairosvg.svg2png(bytestring=svg,write_to=fout)

        fout.close()
        
        
        document = pyDocx()
        document.add_heading('Patient treatment report')
        
        document.add_page_break()
        document.add_heading('Treatment description:', level=2)
        
        document.add_paragraph(t_treatment)
        document.add_paragraph(t_n_patients)
        document.add_paragraph(t_hist)
        document.add_paragraph(t_grade)
        document.add_paragraph(t_hormone)
        document.add_paragraph(t_her2)
        document.add_paragraph(t_stage)
        document.add_paragraph(t_cit)
        document.add_paragraph(o_org)       
        document.add_picture(settings.MEDIA_ROOT+'/output.png', width=Inches(7.0))
        
        document.add_page_break()
        
        f = StringIO()
        document.save(f)
        
        response = HttpResponse(f.getvalue(), mimetype='application/vnd.ms-word')
        response['Content-Disposition'] = 'attachment; filename=PatientReport.docx'
        return response 
        


    
    def render_to_json_response(self, context, **response_kwargs):
        data = json.dumps(context)
        response_kwargs['content_type'] = 'application/json'
        return HttpResponse(data, **response_kwargs)
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(MedicAjaxGenerateReport, self).dispatch(request, *args, **kwargs)
    
    def get_context_data(self, **kwargs):
              
        context = super(MedicAjaxGenerateReport, self).get_context_data(**kwargs)
        
        
        context['test'] = "TEst"
        
        return context
    
       
class MedicAjaxGenerateFullReport(TemplateView):
    template_name = 'medic/patient_treatment_pdf.html'
    
    def post(self, request, *args, **kwargs):
        
        
        document = pyDocx()
        document.add_heading('Patient treatment report')
        
        table = document.add_table(1, 2)
        table.allow_autofit = True
        table.style = 'TableGrid'
        # populate header row --------
        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Treatment'
        heading_cells[1].text = 'Patient status'
        
        
        
        
        for treat in TreatmentMethod.objects.filter(nosology = 1):
            treat_id = str(treat.id)
            cells = table.add_row().cells
            cells[0].text = request.POST.get('t_treatment'+treat_id)
            cells[1].text = request.POST.get('t_responder'+treat_id)
            
        document.add_page_break()
        
        for treat in TreatmentMethod.objects.filter(nosology = 1):
            
            treat_id = str(treat.id)
            
            t_n_patients = request.POST.get('t_n_patients'+treat_id)
            t_hist = request.POST.get('t_hist'+treat_id)
            t_grade = request.POST.get('t_grade'+treat_id)
            t_hormone = request.POST.get('t_hormone'+treat_id)
            t_her2 = request.POST.get('t_her2'+treat_id)
            t_stage = request.POST.get('t_stage'+treat_id)
            t_treatment = request.POST.get('t_treatment'+treat_id)
            t_cit = request.POST.get('t_cit'+treat_id)
            t_org = request.POST.get('t_org'+treat_id)
            svg = request.POST.get('svg'+treat_id)
            
            fout = open(settings.MEDIA_ROOT+'/medic/output.png','w')
        
            cairosvg.svg2png(bytestring=svg,write_to=fout)

            fout.close()
            
            document.add_heading('Treatment description:', level=2)
        
            document.add_paragraph('Treatment: '+t_treatment)
            document.add_paragraph('Number of patients: '+t_n_patients)
            document.add_paragraph('Histological type: '+t_hist)
            document.add_paragraph('Grade: ' + t_grade)
            document.add_paragraph('Hormone receptor status: '+t_hormone)
            document.add_paragraph('HER2 status: '+t_her2)
            document.add_paragraph('Stage: '+t_stage)
            document.add_paragraph('Citation(s): '+t_cit)
            document.add_paragraph('Organization name: '+t_org)       
            document.add_picture(settings.MEDIA_ROOT+'/medic/output.png', width=Inches(7.0))
        
            document.add_page_break()
            
        f = StringIO()
        document.save(f)
        
        response = HttpResponse(f.getvalue(), mimetype='application/vnd.ms-word')
        response['Content-Disposition'] = 'attachment; filename=PatientReport.docx'
        return response 
            
            
    
    def render_to_json_response(self, context, **response_kwargs):
        data = json.dumps(context)
        response_kwargs['content_type'] = 'application/json'
        return HttpResponse(data, **response_kwargs)
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(MedicAjaxGenerateFullReport, self).dispatch(request, *args, **kwargs)
    
    def get_context_data(self, **kwargs):
              
        context = super(MedicAjaxGenerateFullReport, self).get_context_data(**kwargs)
        
        
        context['test'] = "TEst"
        
        return context    
    
    
    
    