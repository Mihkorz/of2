# -*- coding: utf-8 -*-
import json
import pandas as pd
import numpy as np
import networkx as nx
import itertools
import csv
import random 

from django.views.generic.base import TemplateView
from django.views.generic.detail import DetailView
from django.views.generic import ListView
from django.shortcuts import redirect
from django.contrib.auth import logout
from django.http import HttpResponse
from django.conf import settings
from django.utils.decorators import method_decorator
from django.contrib.auth.decorators import login_required
from django.contrib.admin.views.decorators import staff_member_required

from .models import Report, GeneGroup, PathwayGroup, TfGroup
from core.models import Pathway, Node, Component

class ReportList(ListView):
    model = Report
    template_name = 'report/report_list.html'
    context_object_name = 'reports'
    paginate_by = 100
    
    @method_decorator(staff_member_required)
    def dispatch(self, request, *args, **kwargs):
        return super(ReportList, self).dispatch(request, *args, **kwargs)
    
    def get_context_data(self, **kwargs):
              
        context = super(ReportList, self).get_context_data(**kwargs)
        
        
        context['test'] = "Test"
        
        return context

class ReportDetail(DetailView):
    model = Report
    template_name = 'report/report_permition_denied.html'#'report/report_detail.html'
    
    def is_member(self, user, group_name):
        return user.groups.filter(name=group_name).exists()
    
    @method_decorator(login_required)
    def dispatch(self, request, *args, **kwargs):
        
        user = request.user
        
        #standard
        if  (user.is_staff):
            self.template_name = 'report/report_detail_abovebeyond.html'
        
        
        if (self.get_object().id==4 or self.get_object().id==8) and (self.is_member(user, 'Astrazeneca') or user.is_staff):
            self.template_name = 'report/report_detail_az.html' 
        if 'jjms' in self.get_object().slug and (self.is_member(user, 'J&J') or user.is_staff):
            self.template_name = 'report/report_detail_jjms.html'
        if ('above-and-beyond' in self.get_object().slug ) and (self.is_member(user, 'Above and Beyond') or user.is_staff):
            self.template_name = 'report/report_detail_abovebeyond.html'
        if 'Hnkl' in self.get_object().slug and (self.is_member(user, 'Henkel') or user.is_staff):
            self.template_name = 'report/report_detail_henkel.html'
        if 'GSK-NCGOM-2016N298312' in self.get_object().slug and (self.is_member(user, 'GSK') or user.is_staff):
            self.template_name = 'report/report_detail_GSK-NCGOM.html'
        if 'gsk' in self.get_object().slug and (self.is_member(user, 'GSK') or user.is_staff):            
            self.template_name = 'report/report_detail_abovebeyond.html'
        if 'gsk_prj2_' in self.get_object().slug and (self.is_member(user, 'GSK') or user.is_staff):          
            self.template_name = 'report/gsk_prj2_1d.html'
            if '4d_high' in self.get_object().slug:
                self.template_name = 'report/gsk_prj2_4d_high.html'
            if '4d_medium_low' in self.get_object().slug:
                self.template_name = 'report/gsk_prj2_4d_medium_low.html'
        if 'nova' in self.get_object().slug and (self.is_member(user, 'Novartis') or user.is_staff):
            self.template_name = 'report/report_detail_novartis_prj1.html'
            
        if 'nova_prj1_biochem' in self.get_object().slug and (self.is_member(user, 'Novartis') or user.is_staff):
            self.template_name = 'report/just_text.html'
            
        if 'inswx-report' in self.get_object().slug and (self.is_member(user, 'INSWX') or user.is_staff):
            self.template_name = 'report/inswx.html'
        
                
                
        
        
        #user.is_staff
        #raise Exception('stop')
        return super(ReportDetail, self).dispatch(request, *args, **kwargs)
    
    def get_context_data(self, **kwargs):
              
        context = super(ReportDetail, self).get_context_data(**kwargs)
        
        """
        df_logfc = pd.read_csv('/home/mikhail/Downloads/Aliper/DEG-OUTPUT-PROPER/expression-data-Doxorubicin-IPANDA.tsv_DEG.txt', sep='\t',
                                 index_col='SYMBOL')
        
        df_logfc.to_csv('/home/mikhail/Downloads/Aliper/Doxorubicin-IPANDA.tsv_DEG.txt', sep=',')
        raise Exception('stop')
        """
        context['test'] = "Test"
        
        return context

class ReportGeneVolcanoJson(TemplateView):
    template_name="website/bt_report.html"
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(ReportGeneVolcanoJson, self).dispatch(request, *args, **kwargs)
    
    def post(self, request, *args, **kwargs):
        
        json_data = json.loads(request.body)
        
        file_name = json_data['file_name']
         
        report = Report.objects.get(id=json_data['reportID'])
        pval_tres = report.pval_theshold_plot
        logFC_tres = report.logcf_theshold_plot
        

        try:
            df_gene = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name)
            df_gene = df_gene[['gene', 'logFC', 'adj.P.Val', 'P.Value']]
            #df_gene.rename(columns={'gene': 'Symbol'}, inplace=True)
        except:
            try:
                df_gene = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name, sep=None)
                df_gene = df_gene[['SYMBOL', 'logFC', 'adj.P.Val', 'P.Value']]
            except:
                try:
                    df_gene = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name, sep=None)
                    df_gene = df_gene[['SYMBOL', 'logFC', 'adj.P.Val']]
                except:
                    try:
                        df_gene = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name, sep=None)
                        df_gene = df_gene[['SYMBOL', 'log2FoldChange', 'padj']]
                        df_gene.rename(columns={'log2FoldChange': 'logFC', 'padj': 'adj.P.Val'}, inplace=True)
                    except: #For Henkel Report ONLY!
                        df_gene = pd.read_csv(settings.MEDIA_ROOT+"/report-portal/"+report.slug+"/Henkel_mean.csv", sep=None)
                        df_gene = df_gene[['SYMBOL', 'logFC', 'adj.P.Val']]
                
        df_gene.fillna(0, inplace=True)
                        
        df_gene['logFC'] = df_gene['logFC'].round(decimals=2)
        
        df_gene['_row'] = df_gene['adj.P.Val'].round(decimals=2)
        
        df_gene_copy = df_gene.copy()  
         
        if 'nova_' in report.slug: # for Novartis reports filter only on p-value column
            df_gene = df_gene[(df_gene['P.Value']<pval_tres) & (np.absolute(df_gene['logFC'])>logFC_tres)]
            df_gene['adj.P.Val'] = df_gene['P.Value']
        else:
            df_gene = df_gene[(df_gene['adj.P.Val']<pval_tres) & (np.absolute(df_gene['logFC'])>logFC_tres)]
        
        if df_gene.empty:
                df_gene = df_gene_copy[(df_gene_copy['P.Value']<pval_tres) & (np.absolute(df_gene_copy['logFC'])>logFC_tres)]    
                df_gene['adj.P.Val'] = df_gene['P.Value'] 
         
        df_gene['adj.P.Val'] = -1*np.log10(df_gene['adj.P.Val'])        
         
        #raise Exception("reportID: "+str(report.id))
        output_json = df_gene.to_json(orient='records')        
        
        response_data =json.loads(output_json)
        return HttpResponse(json.dumps(response_data), content_type="application/json")
    
    def get(self, request, *args, **kwargs):
        
        raise Exception('See POST request!')      
                
    
class ReportGeneTableJson(TemplateView): 
    template_name="report/report_detail.html"
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(ReportGeneTableJson, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        
        file_name = request.GET.get('file_name')
        
        report = Report.objects.get(pk=request.GET.get('reportID'))
        pval_tres = report.pval_theshold_plot
        logFC_tres = report.logcf_theshold_plot
        
        if file_name!='all':
            
            try:
                df_gene = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name)
                df_gene = df_gene[['gene', 'logFC', 'adj.P.Val', 'P.Value']]
                #df_gene.rename(columns={'gene': 'Symbol'}, inplace=True)
            except:
                try:
                    df_gene = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name, sep=None)
                    df_gene = df_gene[['SYMBOL', 'logFC', 'adj.P.Val', 'P.Value']]
                except:
                    try:
                        df_gene = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name, sep=None)
                        df_gene = df_gene[['SYMBOL', 'logFC', 'adj.P.Val']]
                    except:
                        try:
                            df_gene = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name, sep=None)
                            df_gene = df_gene[['SYMBOL', 'log2FoldChange', 'padj']]
                            df_gene.rename(columns={'log2FoldChange': 'logFC', 'padj': 'adj.P.Val'}, inplace=True)
                        except: #For Henkel Report ONLY!
                            df_gene = pd.read_csv(settings.MEDIA_ROOT+"/report-portal/"+report.slug+"/Henkel_mean.csv", sep=None)
                            df_gene = df_gene[['SYMBOL', 'logFC', 'adj.P.Val']]
                            
                
                
                df_gene.fillna(1, inplace=True)
                
                
                
            df_gene['logFC'] = df_gene['logFC'].round(decimals=2)
            
            
            df_gene_copy = df_gene.copy()        
        
            
            
            if 'nova_' in report.slug: # for Novartis reports filter only on p-value column
                df_gene = df_gene[(df_gene['P.Value']<pval_tres) & (np.absolute(df_gene['logFC'])>logFC_tres)]
                df_gene['adj.P.Val'] = df_gene['P.Value']
            else:
                df_gene = df_gene[(df_gene['adj.P.Val']<pval_tres) & (np.absolute(df_gene['logFC'])>logFC_tres)]
                
                
            
            if df_gene.empty:
                if "NBA" in file_name:
                    df_gene = df_gene_copy[(np.absolute(df_gene_copy['logFC'])>logFC_tres)]
                else:    
                    df_gene = df_gene_copy[(df_gene_copy['P.Value']<pval_tres) & (np.absolute(df_gene_copy['logFC'])>logFC_tres)]    
                df_gene['adj.P.Val'] = df_gene['P.Value']
            
            #raise Exception('stop')            
                        
            df_gene['adj.P.Val'] = df_gene['adj.P.Val'].map('{:,.2e}'.format)
            df_gene['adj.P.Val'] = df_gene['adj.P.Val'].apply(str) 
        
    
        else: #file_name=='all'
            
            lgroups = []
            for group in report.genegroup_set.all():
                
                try:
                    df = pd.read_csv(group.doc_logfc.path, sep=None, index_col='SYMBOL')
                    df = df[['logFC', 'adj.P.Val']]
                except:
                    df = pd.read_csv(group.doc_logfc.path, sep=None, index_col='SYMBOL')
                    df = df[['log2FoldChange', 'padj']]
                    df.rename(columns={'log2FoldChange': 'logFC', 'padj': 'adj.P.Val'}, inplace=True)
                
                #df.rename(columns={'log2FoldChange': 'logFC'}, inplace=True)
                
                #raise Exception('stop')
                
                lgroups.append(df)
            
            df_gene = pd.DataFrame()
            
            for idx, val in enumerate(lgroups):
                df_gene[idx] = val['logFC'].round(decimals=2)

            df_gene.reset_index(inplace=True)
            #raise Exception('gene table')

        if 'Hnkl' in report.slug: #for Henkel report only!!!
            #pass
            df_gene = df_gene.astype(str)
            df_gene = df_gene.apply(lambda x: x.str.replace(".",","))
            
             
        
        output_json = df_gene.to_json(orient='values')
        response_data = {'data': json.loads(output_json)}
        
        return HttpResponse(json.dumps(response_data), content_type="application/json")

class ReportGeneScatterJson(TemplateView):
    template_name="report/report_detail.html"
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(ReportGeneScatterJson, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        
        file_name = settings.MEDIA_ROOT+'/'+request.GET.get('file_name')
        report = Report.objects.get(pk=request.GET.get('reportID'))
        logFC_tres = report.logcf_theshold_plot
        pval_tres = report.pval_theshold_plot
        
        with open(file_name, 'rb') as csvfile:
            sniffer = csv.Sniffer()
            dialect = sniffer.sniff(csvfile.read(), delimiters='\t,;')
            csvfile.seek(0)
            
        df_gene = pd.read_csv(file_name, delimiter=dialect.delimiter,
                                 index_col='SYMBOL')
        
        
        df_tumour = df_gene[[x for x in df_gene.columns if 'Tumour' in x]]
        s_tumour = df_tumour.mean(axis=1).round(decimals=2)
        
        df_norm = df_gene[[x for x in df_gene.columns if 'Norm' in x]]
        s_norm = df_norm.mean(axis=1).round(decimals=2)        
        
        df_output = pd.DataFrame()
        
        df_output['x'] = s_tumour
        df_output['y'] = s_norm
        
        df_output['FC'] = s_tumour.divide(s_norm)
        
        df_output = np.log2(df_output)
        try:
            df_output['adj.P.Val'] = df_gene['adj.P.Val']
            df_output = df_output[df_output['adj.P.Val']<pval_tres]
        except:
            pass
        
        df_output = df_output[np.absolute(df_output['FC'])>logFC_tres]

        #raise Exception('scatter')
        df_output = df_output[df_output['x']>0 ]
        df_output = df_output[df_output['y']>0 ]
        
        df_output.index.name='name'
        df_output.reset_index(inplace=True)
        
        df_output = df_output.to_json(orient='records')        
        
        
        response_data =  json.loads(df_output)
        return HttpResponse(json.dumps(response_data), content_type="application/json")
        
class ReportGeneTableScatterJson(TemplateView): 
    template_name="report/report_detail.html"
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(ReportGeneTableScatterJson, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        
        file_name = request.GET.get('file_name')
        report = Report.objects.get(pk=request.GET.get('reportID'))
        logFC_tres = report.logcf_theshold_plot
        pval_tres = report.pval_theshold_plot
        
        if file_name!='all':          
            
            df_gene = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name)
            try:
                df_gene = df_gene[['gene', 'logFC', 'adj.P.Val']]
            except:
                df_gene = df_gene[['SYMBOL', 'logFC', 'adj.P.Val']]
                
            df_gene['logFC'] = df_gene['logFC'].round(decimals=2)
                     
            df_gene.replace([np.inf, -np.inf], np.nan, inplace=True)
            df_gene = df_gene[(np.absolute(df_gene['logFC'])>logFC_tres)] 
            try:
                df_gene = df_gene[df_gene['adj.P.Val']<pval_tres]
            except:
                pass      
            
        
        else:
            report = Report.objects.get(pk=request.GET.get('reportID'))
            lgroups = []
            for group in report.genegroup_set.all():
                lgroups.append(pd.read_csv(group.doc_logfc.path, index_col='SYMBOL'))
            
            df_gene = pd.DataFrame()
            
            for idx, val in enumerate(lgroups):
                df_gene[idx] = val['logFC'].round(decimals=2)

            df_gene.reset_index(inplace=True)
            raise Exception('gene table')

        output_json = df_gene.to_json(orient='values')
        response_data = {'data': json.loads(output_json)}
        
        return HttpResponse(json.dumps(response_data), content_type="application/json")        
    
    
        
class ReportGeneBoxplotJson(TemplateView):
    
    template_name="report/report_detail.html"
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(ReportGeneBoxplotJson, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        
        gene  = request.GET.get('gene')        
        report = Report.objects.get(pk=request.GET.get('reportID'))
        categories = request.GET.get('categories')
        
         
        
        df_list = []
        
        try: #old report with precalculated box plot file
            for group in report.genegroup_set.all():
            
                df_group = pd.read_csv(group.doc_boxplot.path,
                                   index_col='SYMBOL', sep='\t')
                df_list.append(df_group)
        
            series_tumour = []
            series_norm = []
            i=0
            for df in df_list:          
                
                filered_df = df[[x for x in df if 'Tumour' in x]]
                
                row_gene = filered_df.loc[gene]      
        
                median = np.around(np.log2(row_gene['Tumour_median']), decimals=2)  
                upper_quartile = np.around(np.log2(row_gene['Tumour_upper_quartile']), decimals=2) 
                lower_quartile = np.around(np.log2(row_gene['Tumour_lower_quartile'] ), decimals=2)           
                upper_whisker = np.around(np.log2(row_gene['Tumour_upper_whisker']), decimals=2) 
                lower_whisker = np.around(np.log2(row_gene['Tumour_lower_whisker']), decimals=2) 
              
                lSerie = [lower_whisker, lower_quartile, median, upper_quartile, upper_whisker]
                 
                series_tumour.append(lSerie)
            
                if i==0:
                    pass
                """
                    filered_df = df[[x for x in df if 'Norm' in x]]
                
                    row_gene = filered_df.loc[gene]      
        
                    median = np.around(np.log2(row_gene['Normal_median']), decimals=2) 
                    upper_quartile = np.around(np.log2(row_gene['Normal_upper_quartile']), decimals=2) 
                    lower_quartile = np.around(np.log2(row_gene['Normal_lower_quartile']), decimals=2)            
                    upper_whisker = np.around(np.log2(row_gene['Normal_upper_whisker']), decimals=2) 
                    lower_whisker = np.around(np.log2(row_gene['Normal_lower_whisker']), decimals=2)
              
                    lSerie = [lower_whisker, lower_quartile, median, upper_quartile, upper_whisker]
                    series_tumour.append(lSerie)                
                """
                i=i+1
        except: # AZ report style
            
            lcategories = categories.split(',')
            
            
                
            if len(lcategories)<2: # if there is only one category. 4 volcano table only
                groups = GeneGroup.objects.filter(name=lcategories[0], report=report)
                 
            else:
                groups = report.genegroup_set.all()
                
                       
            
            for group in groups:
                try:
                    df_group = pd.read_csv(group.document.path,
                                   index_col='SYMBOL', sep=None)
                except:
                    df_group = pd.read_csv(group.document.path,
                                   index_col='SYMBOL', sep='\t')
                df_list.append(df_group.loc[gene])
            
            series_tumour = []
            
            i=0
            for df in df_list:
                
                r_tumour = df.filter(like='Tumour')                
                
                median = np.median(r_tumour)               
                upper_quartile = np.percentile(r_tumour, 75)
                lower_quartile = np.percentile(r_tumour, 25)
                iqr = upper_quartile - lower_quartile
                upper_whisker = r_tumour[r_tumour<=upper_quartile+1.5*iqr].max()
                lower_whisker = r_tumour[r_tumour>=lower_quartile-1.5*iqr].min()
                
                # Now we take log2 if needed
                apply_log = True
                if 'gsk_prj2_' in report.slug: # expression data already with log2 applied
                    apply_log = False
                    
                if apply_log:
                    median = np.log2(median)
                    upper_quartile = np.log2(upper_quartile)
                    lower_quartile = np.log2(lower_quartile)
                    upper_whisker = np.log2(upper_whisker)
                    lower_whisker = np.log2(lower_whisker)
        
                 
                if np.isnan(lower_whisker) or np.isinf(lower_whisker):
                    lower_whisker = lower_quartile
                if np.isnan(upper_whisker) or np.isinf(upper_whisker):
                    upper_whisker = upper_quartile
              
                lSerie = [lower_whisker, lower_quartile, median, upper_quartile, upper_whisker]
                
            
                #lSerie[lSerie == -np.inf] = 0
                
                
                #lSerie.replace([np.inf, -np.inf], 0)
                 
                series_tumour.append(lSerie)
                    
             
            #Get Norms boxplot
            ldf_norm = []
            if report.id==8 and report.id==6:
                group_a549 = report.genegroup_set.filter(name__icontains='A549')[0] 
                df_norm_a549 = pd.read_csv(group_a549.document.path,
                                   index_col='SYMBOL', sep='\t')
                group_mcf7 = report.genegroup_set.filter(name__icontains='MCF7')[0]
                df_norm_mcf7 = pd.read_csv(group_mcf7.document.path,
                                   index_col='SYMBOL', sep='\t')
            
                if len(lcategories)<2:
                    if 'A549' in lcategories[0]:
                        ldf_norm.append(df_norm_a549.loc[gene])
                    if 'MCF7' in lcategories[0]:
                        ldf_norm.append(df_norm_mcf7.loc[gene])
                else:
                    ldf_norm.append(df_norm_a549.loc[gene])
                    ldf_norm.append(df_norm_mcf7.loc[gene])
            else:
                ldf_norm.append(df_group.loc[gene])
                
                           
              
            #raise Exception('haha') 29
            
            if report.id == 29: # Novartis Nasal
                df_norm = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/nova_prj1/nasal_norm.csv', index_col='SYMBOL')
                
                ldf_norm = [df_norm.loc[gene]]
                
            if report.id == 31: # Novartis Nasal
                df_norm = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/nova_prj1_pbmc/pbmc_norm.csv', index_col='SYMBOL')
                
                ldf_norm = [df_norm.loc[gene]]
                
            for df_norm in ldf_norm:       
                
                
                r_norm = df_norm.filter(like='Norm')
                                
                median = np.median(r_norm)
                upper_quartile = np.percentile(r_norm, 75)
                lower_quartile = np.percentile(r_norm, 25)
                iqr = upper_quartile - lower_quartile
                upper_whisker = r_norm[r_norm<=upper_quartile+1.5*iqr].max()
                lower_whisker = r_norm[r_norm>=lower_quartile-1.5*iqr].min()
                
                if apply_log:
                    median = np.log2(median)
                    upper_quartile = np.log2(upper_quartile)
                    lower_quartile = np.log2(lower_quartile)
                    upper_whisker = np.log2(upper_whisker)
                    lower_whisker = np.log2(lower_whisker)
                
                if np.isnan(lower_whisker) or np.isinf(lower_whisker):
                    lower_whisker = lower_quartile
                if np.isnan(upper_whisker) or np.isinf(upper_whisker):
                    upper_whisker = upper_quartile
                
                lNorm = [lower_whisker, lower_quartile, median, upper_quartile, upper_whisker]
                
                #raise Exception('stop')
                if report.id not in [6, 35, 32]:
                    series_tumour.append(lNorm)
                else:
                    if len(lcategories)<2:
                        series_tumour.append(lNorm)
  
            
            #raise Exception('boxplot') 
        s1 = {
              'name': 'boxplot',              
              'data': series_tumour,
              'tooltip': {
                          'headerFormat': '<em>Group: {point.key}</em><br/>'
                          }
              }        
        
        
        response_data = s1
        #raise Exception('boxplot') 
        return HttpResponse(json.dumps(response_data), content_type="application/json")  

class ReportGeneBarplotJson(TemplateView):
    template_name="website/report.html"
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(ReportGeneBarplotJson, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        
        gene  = request.GET.get('gene')        
        report = Report.objects.get(pk=request.GET.get('reportID'))
        lcategories = [x.strip() for x in request.GET.get('categories').split(',')]
        
        
        
        df_list = []
        cat_names = []
        
        if len(lcategories)<2: # if there is only one category. 4 volcano table only
            groups = GeneGroup.objects.filter(name=lcategories[0], report=report)
                 
        else:
            groups = report.genegroup_set.all()
        
        
        if lcategories == ['henkel_all']:
            groups = ['henkel_all']
            
        for group in groups:
            try:
                df_group = pd.read_csv(group.document.path,
                                index_col='SYMBOL', sep=None)
            except:
                try:
                    df_group = pd.read_csv(group.document.path,
                                   index_col='SYMBOL', sep='\t')
                except: #For Henkel Report ONLY!\
                    df_group = pd.read_csv(settings.MEDIA_ROOT+"/report-portal/"+report.slug+"/Henkel_mean_exp.csv", 
                                           index_col='SYMBOL', sep=None)
                    
    
            df_list.append(df_group.loc[gene])
            
            if group == 'henkel_all':
                cat_names.append('Poolw')
            else:
                cat_names.append(group.name)
            
            
            series_tumour = []
            series_norm = []
            
            i=0
            for df in df_list:
                
                r_tumour = df.filter(like='Tumour')
                r_norm = df.filter(like='Norm')
                
                series_tumour.append(r_tumour[0])
                series_norm.append(r_norm[0])
        
        
        response_data = {
                         'name': 'barplot',
                         'categories_name': cat_names,
                         'tumour': [ round(elem, 2) for elem in series_tumour ],
                         'norm': [ round(elem, 2) for elem in series_norm ]
                         }  
        
        #raise Exception('barplot!!!')
        return HttpResponse(json.dumps(response_data), content_type="application/json")    
        
        
        
    
class ReportPathwayTableJson(TemplateView):
    template_name="website/report.html"
    def dispatch(self, request, *args, **kwargs):
        
        return super(ReportPathwayTableJson, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        
        file_name1 = request.GET.get('file_name1')
        file_name2 = request.GET.get('file_name2')
        is_metabolic = request.GET.get('is_metabolic')
        
        report = Report.objects.get(pk=request.GET.get('reportID'))
        
        if is_metabolic =='false':
            is_metabolic = False
        else:
            is_metabolic = True
        
        if file_name1 == file_name2 == 'all':
            
            lgroups = []
            for group in report.pathwaygroup_set.all():
                df_path = pd.read_csv(group.doc_proc.path, index_col='Pathway')
                
                if is_metabolic:
                    df_path = df_path[df_path['Database']=='metabolism']
                else:
                    df_path = df_path[df_path['Database']!='metabolism']
                    #df_path = df_path[df_path['Database']!='kegg']
                    
                     
                lgroups.append(df_path)
            
            if 'Hnkl' in report.slug: # for Henkel reports only!!!
                df_mean = pd.DataFrame(index=df_path.index)
                for df_group in lgroups:
                    df_mean = df_mean.join(df_group['0'], how='inner', lsuffix='_1')                    
                
                df_mean = pd.DataFrame(df_mean.mean(axis=1))
                df_mean.columns = ['0']
                lgroups.insert(0, df_mean)
                
            
            df_output = pd.DataFrame()
            
            for idx, val in enumerate(lgroups):
                df_output[idx] = val['0'].round(decimals=2)
                if 'inswx-report' in report.slug:
                    df_output[str(idx)+'_pval'] = val['p_ora']
                    #raise Exception('output')
            
            try:
                df_output.drop(['Target_drugs_pathway'], inplace=True)
            except:
                pass
            df_output.reset_index(inplace=True)
            
            #raise Exception('haha')            
            
        else:
            
            df_1 = pd.read_excel(settings.MEDIA_ROOT+"/../static/report/loreal/"+file_name1,
                                sheetname='PAS1', index_col='Pathway')
            df_2 = pd.read_excel(settings.MEDIA_ROOT+"/../static/report/loreal/"+file_name2,
                                 sheetname='PAS1', index_col='Pathway')
        
            
            if is_metabolic:
                df_1 = df_1[df_1['Database']=='metabolism']
                df_2 = df_2[df_2['Database']=='metabolism']
            else:
                df_1 = df_1[df_1['Database']!='metabolism']
                #df_1 = df_1[df_1['Database']!='kegg']
                df_2 = df_2[df_2['Database']!='metabolism']
                #df_2 = df_2[df_2['Database']!='kegg']
            
            df1_tumour = df_1[[x for x in df_1.columns if 'Tumour' in x]]
            s1_tumour = df1_tumour.mean(axis=1).round(decimals=2)
        
            df2_tumour = df_2[[x for x in df_2.columns if 'Tumour' in x]]
            s2_tumour = df2_tumour.mean(axis=1).round(decimals=2)               
        
            df_output = pd.DataFrame()
        
            df_output['1'] = s1_tumour
            df_output['2'] = s2_tumour
            try:
                df_output.drop(['Target_drugs_pathway'], inplace=True)
            except:
                pass
            df_output.reset_index(inplace=True)
        
        df_output.fillna(0, inplace=True)
        
        if 'Hnkl' in report.slug: #for Henkel report only!!!
            df_output = df_output.astype(str)
            df_output = df_output.apply(lambda x: x.str.replace(".",","))
        
        df_json = df_output.to_json(orient='values')
        
        
        #raise Exception('path table')
        
        response_data = {'aaData': json.loads(df_json)}
        return HttpResponse(json.dumps(response_data), content_type="application/json")
    
class ReportAjaxPathwayVenn(TemplateView):
    template_name="website/report_ajax_venn.html"
    def dispatch(self, request, *args, **kwargs):
        return super(ReportAjaxPathwayVenn, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
          
        report = Report.objects.get(pk=request.GET.get('reportID'))
        pval_tres = report.pval_theshold_compare
        logFC_tres = report.logcf_theshold_compare
        pas_tres = report.pas_theshold
        
        group_names = self.request.GET.get('group_names')
        is_metabolic = self.request.GET.get('is_metabolic')
        regulation = self.request.GET.get('regulation')
        path_gene = self.request.GET.get('path_gene')
        
        lgroup = group_names.split('vs')
        
        name1 = lgroup[0]
        name2 = lgroup[1]
        try:
            name3 = lgroup[2]
            compare3 = True
        except:
            compare3 = False
        try:
            name4 = lgroup[3]
            compare4 = True
        except:
            compare4 = False
        #raise Exception('VENN!!!!')
        
        venn_cyrcles = []
        
        
        if path_gene == 'pathways':
            
            group1 = PathwayGroup.objects.get(name=lgroup[0], report=report)
            file_name1 = group1.doc_proc
            df1 = pd.read_csv(file_name1, index_col='Pathway')
            df1 = df1[np.absolute(df1['0'])>pas_tres]
            if 'inswx-report' in report.slug:
                df1 = df1[np.absolute(df1['p_ora'])<0.05]
                
            
            group2 = PathwayGroup.objects.get(name=lgroup[1], report=report)
            file_name2 = group2.doc_proc
            df2 = pd.read_csv(file_name2, index_col='Pathway')
            df2 = df2[np.absolute(df2['0'])>pas_tres]
            if 'inswx-report' in report.slug:
                df2 = df2[np.absolute(df2['p_ora'])<0.05]
            
            if compare3:
                group3 = PathwayGroup.objects.get(name=lgroup[2], report=report)
                file_name3 = group3.doc_proc
                df3 = pd.read_csv(file_name3, index_col='Pathway')
                df3 = df3[np.absolute(df3['0'])>pas_tres]
                if 'inswx-report' in report.slug:
                    df3 = df3[np.absolute(df3['p_ora'])<0.05]          
            
            if compare4:
                group4 = PathwayGroup.objects.get(name=lgroup[3], report=report)
                file_name4 = group4.doc_proc
                df4 = pd.read_csv(file_name4, index_col='Pathway')
                df4 = df4[np.absolute(df4['0'])>pas_tres]
                if 'inswx-report' in report.slug:
                    df4 = df4[np.absolute(df4['p_ora'])<0.05] 
                
                                    
            
            #raise Exception('VENN!!!!')           
            
            if is_metabolic=='true':
                df1 = df1[df1['Database']=='metabolism']
                df2 = df2[df2['Database']=='metabolism']
                if compare3:
                    df3 = df3[df3['Database']=='metabolism']
                if compare4:
                    df4 = df4[df4['Database']=='metabolism']
            else:
                df1 = df1[df1['Database']!='metabolism']
                df2 = df2[df2['Database']!='metabolism']
                if compare3:
                    df3 = df3[df3['Database']!='metabolism']
                if compare4:
                    df4 = df4[df4['Database']!='metabolism']
                
        elif path_gene == 'genes':
            
            group1 = GeneGroup.objects.get(name=lgroup[0], report=report)
            file_name1 = group1.doc_logfc.path
            try:
                df1 = pd.read_csv(file_name1, index_col='SYMBOL', sep=None)
                df1 = df1[['logFC', 'adj.P.Val']]
            except:
                df1 = pd.read_csv(file_name1, index_col='SYMBOL', sep=None)
                df1 = df1[['log2FoldChange', 'padj']]
                df1.rename(columns={'log2FoldChange': 'logFC', 'padj': 'adj.P.Val'}, inplace=True)
            df1.replace([np.inf, -np.inf], 0, inplace=True)
            
            
            if (df1['adj.P.Val']>pval_tres).all():
                df1 = df1[(np.absolute(df1['logFC'])>logFC_tres)]         
            else:
                df1 = df1[(df1['adj.P.Val']<pval_tres) & (np.absolute(df1['logFC'])>logFC_tres)]
                   
            df1 = pd.DataFrame(df1['logFC'])
            df1.columns = ['0'] 
            
            group2 = GeneGroup.objects.get(name=lgroup[1], report=report)
            file_name2 = group2.doc_logfc.path
            try:
                df2 = pd.read_csv(file_name2, index_col='SYMBOL', sep=None)
                df2 = df2[['logFC', 'adj.P.Val']]
            except:
                df2 = pd.read_csv(file_name2, index_col='SYMBOL', sep=None)
                df2 = df2[['log2FoldChange', 'padj']]
                df2.rename(columns={'log2FoldChange': 'logFC', 'padj': 'adj.P.Val'}, inplace=True)
            df2.replace([np.inf, -np.inf], 0, inplace=True)
            
            if (df2['adj.P.Val']>pval_tres).all():
                df2 = df2[(np.absolute(df2['logFC'])>logFC_tres)]
            else:
                df2 = df2[(df2['adj.P.Val']<pval_tres) & (np.absolute(df2['logFC'])>logFC_tres)]
            
            df2 = pd.DataFrame(df2['logFC'])
            df2.columns = ['0']
            
            if compare3:
                group3 = GeneGroup.objects.get(name=lgroup[2], report=report)
                file_name3 = group3.doc_logfc.path
                try:
                    df3 = pd.read_csv(file_name3, index_col='SYMBOL', sep=None)
                    df3 = df3[['logFC', 'adj.P.Val']]
                except:
                    df3 = pd.read_csv(file_name3, index_col='SYMBOL', sep=None)
                    df3 = df3[['log2FoldChange', 'padj']]
                    df3.rename(columns={'log2FoldChange': 'logFC', 'padj': 'adj.P.Val'}, inplace=True)
                df3.replace([np.inf, -np.inf], 0, inplace=True)
                
                if (df3['adj.P.Val']>pval_tres).all():
                    df3 = df3[(np.absolute(df3['logFC'])>logFC_tres)]
                else:
                    df3 = df3[(df3['adj.P.Val']<pval_tres) & (np.absolute(df3['logFC'])>logFC_tres)]
                    
                df3 = pd.DataFrame(df3['logFC'])
                df3.columns = ['0']
                    
            if compare4:
                group4 = GeneGroup.objects.get(name=lgroup[3], report=report)
                file_name4 = group4.doc_logfc.path
                try:
                    df4 = pd.read_csv(file_name4, index_col='SYMBOL', sep=None)
                    df4 = df4[['logFC', 'adj.P.Val']]
                except:
                    df4 = pd.read_csv(file_name4, index_col='SYMBOL', sep=None)
                    df4 = df4[['log2FoldChange', 'padj']]
                    df4.rename(columns={'log2FoldChange': 'logFC', 'padj': 'adj.P.Val'}, inplace=True)
                df4.replace([np.inf, -np.inf], 0, inplace=True)
                if (df4['adj.P.Val']>pval_tres).all():
                    df4 = df4[(np.absolute(df4['logFC'])>logFC_tres)]
                else:
                    df4 = df4[(df4['adj.P.Val']<pval_tres) & (np.absolute(df4['logFC'])>logFC_tres)]
                    
                df4 = pd.DataFrame(df4['logFC'])
                df4.columns = ['0']        
            
        
        elif path_gene == 'deeplearning': # For GSK prj2 only! not dynamic!
            #raise Exception('extended')
            if is_metabolic=='true':
                df1 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/New_top_model_FC_whole.csv',
                                  index_col='probeset')
                df2 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/New_top_model_FC_EC.csv',
                                  index_col='probeset')
                df3 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/New_top_model_FC_SM.csv',
                                  index_col='probeset')
            else:
                df1 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/New_top_model_FC_ext_whole.csv',
                                  index_col='probeset')
                df2 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/New_top_model_FC_ext_EC.csv',
                                  index_col='probeset')
                df3 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/New_top_model_FC_ext_SM.csv',
                                  index_col='probeset')
                #raise Exception('extended')
            
            
            df1 = df1.sort(columns='mean', ascending=False).head(1000)
            df2 = df2.sort(columns='mean', ascending=False).head(1000)
            df3 = df3.sort(columns='mean', ascending=False).head(1000)
            
            df1.rename(columns={'mean': '0'}, inplace=True)
            df2.rename(columns={'mean': '0'}, inplace=True)
            df3.rename(columns={'mean': '0'}, inplace=True)
            #raise Exception('deeplearning')
        
        dict_s = {}
        list_s = [name1, name2]
        if compare3:
            list_s.append(name3) 
        if compare4:
            list_s.append(name4)
        
            
        st1_tumour = df1['0']
        st1_tumour_up = st1_tumour[st1_tumour>0]
        st1_tumour_down = st1_tumour[st1_tumour<0]
        if regulation == 'updown':
            venn_cyrcles.append({'sets': [name1], 'size': (st1_tumour_up.count()+st1_tumour_down.count()),
                                     'id': '1+updown+'+name1+'+'+is_metabolic})
        elif regulation == 'up':
            venn_cyrcles.append({'sets': [name1], 'size': (st1_tumour_up.count()),
                                     'id': '1+up+'+name1+'+'+is_metabolic})
        elif regulation == 'down':
            venn_cyrcles.append({'sets': [name1], 'size': (st1_tumour_down.count()),
                                     'id': '1+down+'+name1+'+'+is_metabolic})
        dict_s[name1] = st1_tumour
        dict_s[name1+' up'] = st1_tumour_up.index
        dict_s[name1+' down'] = st1_tumour_down.index        
          
           
        st2_tumour = df2['0']
        st2_tumour_up = st2_tumour[st2_tumour>0]
        st2_tumour_down = st2_tumour[st2_tumour<0]
        if regulation == 'updown':
            venn_cyrcles.append({'sets': [name2], 'size': (st2_tumour_up.count()+st2_tumour_down.count()),
                                     'id': '1+updown+'+name2+'+'+is_metabolic})
        elif regulation == 'up':
            venn_cyrcles.append({'sets': [name2], 'size': (st2_tumour_up.count()),
                                     'id': '1+up+'+name2+'+'+is_metabolic})
        elif regulation == 'down':
            venn_cyrcles.append({'sets': [name2], 'size': (st2_tumour_down.count()),
                                     'id': '1+down+'+name2+'+'+is_metabolic})    
        dict_s[name2] = st2_tumour
        dict_s[name2+' up'] = st2_tumour_up.index
        dict_s[name2+' down'] = st2_tumour_down.index
        
        if compare3:  
            st3_tumour = df3['0']
            st3_tumour_up = st3_tumour[st3_tumour>0]
            st3_tumour_down = st3_tumour[st3_tumour<0]
            if regulation == 'updown':
                    venn_cyrcles.append({'sets': [name3], 'size': (st3_tumour_up.count()+st3_tumour_down.count()),
                                     'id': '1+updown+'+name3+'+'+is_metabolic})
            elif regulation == 'up':
                    venn_cyrcles.append({'sets': [name3], 'size': (st3_tumour_up.count()),
                                     'id': '1+up+'+name3+'+'+is_metabolic})
            elif regulation == 'down':
                    venn_cyrcles.append({'sets': [name3], 'size': (st3_tumour_down.count()),
                                     'id': '1+down+'+name3+'+'+is_metabolic})
            dict_s[name3] = st3_tumour
            dict_s[name3+' up'] = st3_tumour_up.index
            dict_s[name3+' down'] = st3_tumour_down.index
        
        if compare4:  
            st4_tumour = df4['0']
            st4_tumour_up = st4_tumour[st4_tumour>0]
            st4_tumour_down = st4_tumour[st4_tumour<0]
            if regulation == 'updown':
                    venn_cyrcles.append({'sets': [name4], 'size': (st4_tumour_up.count()+st4_tumour_down.count()),
                                     'id': '1+updown+'+name4+'+'+is_metabolic})
            elif regulation == 'up':
                    venn_cyrcles.append({'sets': [name4], 'size': (st4_tumour_up.count()),
                                     'id': '1+up+'+name4+'+'+is_metabolic})
            elif regulation == 'down':
                    venn_cyrcles.append({'sets': [name4], 'size': (st4_tumour_down.count()),
                                     'id': '1+down+'+name4+'+'+is_metabolic})
            dict_s[name4] = st4_tumour
            dict_s[name4+' up'] = st4_tumour_up.index
            dict_s[name4+' down'] = st4_tumour_down.index
        
        #raise Exception('VENN!!')        
            
        combinations_2= list(itertools.combinations(list_s, 2)) # get all pairs of items
        for idx, combination in enumerate(combinations_2):
                
            index1_up = dict_s[combination[0]+' up']
            index1_down = dict_s[combination[0]+' down']
            index2_up = dict_s[combination[1]+' up']                
            index2_down = dict_s[combination[1]+' down']
            if regulation == 'updown':
                intersection = len(index1_up.intersection(index2_up))+len(index1_down.intersection(index2_down))
                id_x = '2+updown+'+combination[0]+'vs'+ combination[1]+'+'+is_metabolic
            elif regulation == 'up':
                intersection = len(index1_up.intersection(index2_up))
                id_x = '2+up+'+combination[0]+'vs'+ combination[1]+'+'+is_metabolic
            elif regulation == 'down':
                intersection = len(index1_down.intersection(index2_down))
                id_x = '2+down+'+combination[0]+'vs'+ combination[1]+'+'+is_metabolic
            venn_cyrcles.append({'sets': [combination[0], combination[1]],
                                     'size': intersection,
                                     'id': id_x })
            #raise Exception('comb')
        if compare3:    
            combinations_3= list(itertools.combinations(list_s, 3)) # get all triplets of items
            for idx, combination in enumerate(combinations_3):
                
                index1_up = dict_s[combination[0]+' up']
                index1_down = dict_s[combination[0]+' down']
                index2_up = dict_s[combination[1]+' up']                
                index2_down = dict_s[combination[1]+' down']
                index3_up = dict_s[combination[2]+' up']                
                index3_down = dict_s[combination[2]+' down']
                
                inter_up = (index1_up.intersection(index2_up)).intersection(index3_up)
                inter_down = (index1_down.intersection(index2_down)).intersection(index3_down)
                if regulation == 'updown':
                    intersection = len(inter_up)+len(inter_down)
                    id_x = '3+updown+'+combination[0]+'vs'+combination[1]+'vs'+combination[2]+'+'+is_metabolic
                elif regulation == 'up':
                    intersection = len(inter_up)
                    id_x = '3+up+'+combination[0]+'vs'+combination[1]+'vs'+combination[2]+'+'+is_metabolic
                elif regulation == 'down':
                    intersection = len(inter_down)
                    id_x = '3+down+'+combination[0]+'vs'+combination[1]+'vs'+combination[2]+'+'+is_metabolic
                venn_cyrcles.append({'sets': [combination[0], combination[1], combination[2]],
                                     'size': intersection,
                                     'id': id_x})
            
        if compare4:    
            combinations_4= list(itertools.combinations(list_s, 4)) # get all triplets of items
            for idx, combination in enumerate(combinations_4):
                
                index1_up = dict_s[combination[0]+' up']
                index1_down = dict_s[combination[0]+' down']
                index2_up = dict_s[combination[1]+' up']                
                index2_down = dict_s[combination[1]+' down']
                index3_up = dict_s[combination[2]+' up']                
                index3_down = dict_s[combination[2]+' down']
                index4_up = dict_s[combination[3]+' up']                
                index4_down = dict_s[combination[3]+' down']
                
                inter_up = (index1_up.intersection(index2_up)).intersection(index3_up).intersection(index4_up)
                inter_down = (index1_down.intersection(index2_down)).intersection(index3_down).intersection(index4_down)
                if regulation == 'updown':
                    intersection = len(inter_up)+len(inter_down)
                    id_x = '4+updown+'+combination[0]+'vs'+combination[1]+'vs'+combination[2]+'vs'+combination[3]+'+'+is_metabolic
                elif regulation == 'up':
                    intersection = len(inter_up)
                    id_x = '4+up+'+combination[0]+'vs'+combination[1]+'vs'+combination[2]+'vs'+combination[3]+'+'+is_metabolic
                elif regulation == 'down':
                    intersection = len(inter_down)
                    id_x = '4+down+'+combination[0]+'vs'+combination[1]+'vs'+combination[2]+'vs'+combination[3]+'+'+is_metabolic
                venn_cyrcles.append({'sets': [combination[0], combination[1], combination[2], combination[3]],
                                     'size': intersection,
                                     'id': id_x})
 
            #raise Exception('venn all')
        #raise Exception('VENN!!')

        response_data= venn_cyrcles

        return HttpResponse(json.dumps(response_data), content_type="application/json")
    
class ReportAjaxPathwayVennTable(TemplateView):
    template_name="website/report_ajax_venn.html"
    def dispatch(self, request, *args, **kwargs):
        return super(ReportAjaxPathwayVennTable, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        inter_num = int(self.request.GET.get('inter_num'))
        regulation = self.request.GET.get('regulation')
        members = self.request.GET.get('members')
        is_metabolic = self.request.GET.get('is_metabolic')
        path_gene = self.request.GET.get('path_gene')
        
        lMembers = members.split('vs')
        
        report = Report.objects.get(pk=self.request.GET.get('reportID'))
        pval_tres = report.pval_theshold_compare
        logFC_tres = report.logcf_theshold_compare
        pas_tres = report.pas_theshold
        
        if inter_num == 1:
            if path_gene == 'pathways':
                
                group1 = PathwayGroup.objects.get(name=lMembers[0], report=report)
                df_1 = pd.read_csv(group1.doc_proc.path, index_col='Pathway')
                df_1 = df_1[np.absolute(df_1['0'])>pas_tres]
                if 'inswx-report' in report.slug and path_gene == 'pathways':
                    df_1 = df_1[np.absolute(df_1['p_ora'])<0.05]
                    
                    
                                        
                if is_metabolic=='true':
                    df_1 = df_1[df_1['Database']=='metabolism']
                else:                
                    df_1 = df_1[df_1['Database']!='metabolism']
                
            elif path_gene =='genes':
                
                group1 = GeneGroup.objects.get(name=lMembers[0], report=report)
                
                try:
                    df_1 = pd.read_csv(group1.doc_logfc.path,  index_col='SYMBOL', sep=None)
                    df_1 = df_1[['logFC', 'adj.P.Val']]
                except:
                    df_1 = pd.read_csv(group1.doc_logfc.path,  index_col='SYMBOL', sep=None)
                    df_1 = df_1[['log2FoldChange', 'padj']]
                    df_1.rename(columns={'log2FoldChange': 'logFC', 'padj': 'adj.P.Val'}, inplace=True)
                
                if (df_1['adj.P.Val']>pval_tres).all():
                    df_1 = df_1[(np.absolute(df_1['logFC'])>logFC_tres)]
                else:            
                    df_1 = df_1[(df_1['adj.P.Val']<pval_tres) & (np.absolute(df_1['logFC'])>logFC_tres)]
                df_1 = pd.DataFrame(df_1['logFC'])
                df_1.columns = ['0']
            
            elif path_gene =='deeplearning':
                
                if is_metabolic=='true':
                    try:
                        df_1 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/New_top_model_FC_'+lMembers[0]+'.csv',
                                  index_col='probeset')
                        
                    except:
                        df_1 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/New_top_model_FC_whole.csv',
                                  index_col='probeset')
    
                else:
                    try:
                        df_1 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/New_top_model_FC_ext_'+lMembers[0]+'.csv',
                                  index_col='probeset')
                    except:
                        df_1 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/New_top_model_FC_ext_whole.csv',
                                  index_col='probeset')
                        
                df_1 = df_1.sort(columns='mean', ascending=False).head(1000)
                df_1.rename(columns={'mean': '0'}, inplace=True)
                    
                
                
                
            s_tumour = df_1['0']#.round(decimals=2)
            if 'inswx-report' in report.slug and path_gene == 'pathways':
                s_tumour = df_1[['0', 'p_ora']]
                #raise Exception('fuck')
                    
            s_tumour.name = lMembers[0]
            if regulation == 'updown':
                s_tumour = s_tumour[s_tumour!=0]
            elif regulation == 'up':
                if 'inswx-report' in report.slug and path_gene == 'pathways':
                    s_tumour = s_tumour[s_tumour['0']>0]
                else:                
                    s_tumour = s_tumour[s_tumour>0]
            elif regulation == 'down':  
                if 'inswx-report' in report.slug and path_gene == 'pathways':
                    s_tumour = s_tumour[s_tumour['0']<0]
                else:                
                    s_tumour = s_tumour[s_tumour<0]
            
            df_1_tumour = pd.DataFrame(s_tumour)
            if path_gene =='deeplearning':
                df_1_tumour['genes'] = df_1['genes']
                df_1_tumour.reset_index(inplace=True)
                df_1_tumour.set_index('genes', inplace=True)
                #raise Exception('iter1')  
            df_1_tumour.reset_index(inplace=True)
            df_json = df_1_tumour.to_json(orient='values')
            #raise Exception('p-val') 
        
        elif inter_num == 2:
            if path_gene == 'pathways':
                
                group1 = PathwayGroup.objects.get(name=lMembers[0], report=report)
                df_1 = pd.read_csv(group1.doc_proc.path, index_col='Pathway')
                df_1 = df_1[np.absolute(df_1['0'])>pas_tres]
                if 'inswx-report' in report.slug and path_gene == 'pathways' :
                    df_1 = df_1[np.absolute(df_1['p_ora'])<0.05]
                                
                group2 = PathwayGroup.objects.get(name=lMembers[1], report=report)
                df_2 = pd.read_csv(group2.doc_proc.path, index_col='Pathway')
                df_2 = df_2[np.absolute(df_2['0'])>pas_tres]
                if 'inswx-report' in report.slug and path_gene == 'pathways':
                    df_2 = df_2[np.absolute(df_2['p_ora'])<0.05]
            
                if is_metabolic=='true':
                    df_1 = df_1[df_1['Database']=='metabolism']
                    df_2 = df_2[df_2['Database']=='metabolism']
                else:                
                    df_1 = df_1[df_1['Database']!='metabolism']
                    df_2 = df_2[df_2['Database']!='metabolism']
                
            elif path_gene =='genes':
                
                group1 = GeneGroup.objects.get(name=lMembers[0], report=report)
                group2 = GeneGroup.objects.get(name=lMembers[1], report=report)
                
                try:
                    df_1 = pd.read_csv(group1.doc_logfc.path,  index_col='SYMBOL', sep=None)
                    df_1 = df_1[['logFC', 'adj.P.Val']]
                    df_2 = pd.read_csv(group2.doc_logfc.path,  index_col='SYMBOL', sep=None)
                    df_2 = df_2[['logFC', 'adj.P.Val']] 
                except:
                    df_1 = pd.read_csv(group1.doc_logfc.path,  index_col='SYMBOL', sep=None)
                    df_1 = df_1[['log2FoldChange', 'padj']]
                    df_1.rename(columns={'log2FoldChange': 'logFC', 'padj': 'adj.P.Val'}, inplace=True)
                    df_2 = pd.read_csv(group2.doc_logfc.path,  index_col='SYMBOL', sep=None)
                    df_2 = df_2[['log2FoldChange', 'padj']]
                    df_2.rename(columns={'log2FoldChange': 'logFC', 'padj': 'adj.P.Val'}, inplace=True) 
                
                if (df_1['adj.P.Val']>pval_tres).all():
                    df_1 = df_1[(np.absolute(df_1['logFC'])>logFC_tres)]
                else:           
                    df_1 = df_1[(df_1['adj.P.Val']<pval_tres) & (np.absolute(df_1['logFC'])>logFC_tres)]
                if (df_2['adj.P.Val']>pval_tres).all():
                    df_2 = df_2[(np.absolute(df_2['logFC'])>logFC_tres)]
                else:
                    df_2 = df_2[(df_2['adj.P.Val']<pval_tres) & (np.absolute(df_2['logFC'])>logFC_tres)]
                
                df_1 = pd.DataFrame(df_1['logFC'])
                df_2 = pd.DataFrame(df_2['logFC'])
                df_1.columns = ['0']
                df_2.columns = ['0']
            
            elif path_gene =='deeplearning':
                
                if is_metabolic=='true':
                    try:
                        df_1 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/Top_model_FC_'+lMembers[0]+'.csv',
                                  index_col='probes')                        
                    except:
                        df_1 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/Top_model_FC_final.csv',
                                  index_col='probes')
                    try:
                        df_2 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/Top_model_FC_'+lMembers[1]+'.csv',
                                  index_col='probes')                        
                    except:
                        df_2 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/Top_model_FC_final.csv',
                                  index_col='probes')
    
                else:
                    try:
                        df_1 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/Top_model_FC_ext_'+lMembers[0]+'.csv',
                                  index_col='probes')
                    except:
                        df_1 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/Top_model_FC_ext.csv',
                                  index_col='probes')
                    try:
                        df_2 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/Top_model_FC_ext_'+lMembers[1]+'.csv',
                                  index_col='probes')
                    except:
                        df_2 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/Top_model_FC_ext.csv',
                                  index_col='probes')
                        
                df_1 = df_1.sort(columns='mean', ascending=False).head(1000)
                df_1.rename(columns={'mean': '0'}, inplace=True)
                df_2 = df_2.sort(columns='mean', ascending=False).head(1000)
                df_2.rename(columns={'mean': '0'}, inplace=True)
                
                #raise Exception('iter2')
            
            s_tumour1 = df_1['0']#.round(decimals=2)
            if 'inswx-report' in report.slug  and path_gene == 'pathways':
                s_tumour1 = df_1[['0', 'p_ora']]
            s_tumour1.name = lMembers[0]
            
            s_tumour2 = df_2['0']#.round(decimals=2)
            if 'inswx-report' in report.slug  and path_gene == 'pathways':
                s_tumour2 = df_2[['0', 'p_ora']]
            s_tumour2.name = lMembers[1]
            if regulation == 'updown':
                s_tumour1_up = s_tumour1[s_tumour1>0]
                s_tumour1_down = s_tumour1[s_tumour1<0]
                s_tumour2_up = s_tumour2[s_tumour2>0]
                s_tumour2_down = s_tumour2[s_tumour2<0]
                
                df_up = pd.DataFrame(s_tumour1_up).join(pd.DataFrame(s_tumour2_up), how='inner', sort=True)         
                df_down = pd.DataFrame(s_tumour1_down).join(pd.DataFrame(s_tumour2_down), how='inner', sort=True)
                joined_df = df_up.append(df_down)
            elif regulation == 'up':
                if 'inswx-report' in report.slug  and path_gene == 'pathways':
                    s_tumour1 = s_tumour1[s_tumour1['0']>0]
                    s_tumour2 = s_tumour2[s_tumour2['0']>0]
                else:            
                    s_tumour1 = s_tumour1[s_tumour1>0]
                    s_tumour2 = s_tumour2[s_tumour2>0]
                joined_df = pd.DataFrame(s_tumour1).join(pd.DataFrame(s_tumour2), how='inner', lsuffix='_1')
                
                if path_gene =='deeplearning': 
                    joined_df = joined_df.join(df_1['genes'], how='inner')
                    joined_df.reset_index(inplace=True)
                    joined_df.set_index('genes', inplace=True)
                    
                    #raise Exception('iter2s')
                    
                
            elif regulation == 'down':  
                if 'inswx-report' in report.slug  and path_gene == 'pathways':
                    s_tumour1 = s_tumour1[s_tumour1['0']<0]
                    s_tumour2 = s_tumour2[s_tumour2['0']<0]
                else:            
                    s_tumour1 = s_tumour1[s_tumour1<0]
                    s_tumour2 = s_tumour2[s_tumour2<0]
                joined_df = pd.DataFrame(s_tumour1).join(pd.DataFrame(s_tumour2), how='inner', lsuffix='_1')
            
            if path_gene !='deeplearning':
                joined_df = joined_df.groupby(joined_df.index, level=0).mean()
            joined_df.reset_index(inplace=True)
            df_json = joined_df.to_json(orient='values')
            
        elif inter_num == 3:
            if path_gene == 'pathways':
                group1 = PathwayGroup.objects.get(name=lMembers[0], report=report)
                df_1 = pd.read_csv(group1.doc_proc.path, index_col='Pathway')
                df_1 = df_1[np.absolute(df_1['0'])>pas_tres]
                
                group2 = PathwayGroup.objects.get(name=lMembers[1], report=report)
                df_2 = pd.read_csv(group2.doc_proc.path, index_col='Pathway')
                df_2 = df_2[np.absolute(df_2['0'])>pas_tres]
                
                group3 = PathwayGroup.objects.get(name=lMembers[2], report=report)
                df_3 = pd.read_csv(group3.doc_proc.path,  index_col='Pathway')
                df_3 = df_3[np.absolute(df_3['0'])>pas_tres]
            
                if is_metabolic=='true':
                    df_1 = df_1[df_1['Database']=='metabolism']
                    df_2 = df_2[df_2['Database']=='metabolism']
                    df_3 = df_3[df_3['Database']=='metabolism']
                else:                
                    df_1 = df_1[df_1['Database']!='metabolism']
                    df_2 = df_2[df_2['Database']!='metabolism']
                    df_3 = df_3[df_3['Database']!='metabolism']
            
            elif path_gene =='genes':
                group1 = GeneGroup.objects.get(name=lMembers[0], report=report)
                group2 = GeneGroup.objects.get(name=lMembers[1], report=report)
                group3 = GeneGroup.objects.get(name=lMembers[2], report=report)
                
                try:
                    df_1 = pd.read_csv(group1.doc_logfc.path,  index_col='SYMBOL', sep=None)
                    df_2 = pd.read_csv(group2.doc_logfc.path,  index_col='SYMBOL', sep=None) 
                    df_3 = pd.read_csv(group3.doc_logfc.path,  index_col='SYMBOL', sep=None) 
                except:
                    df_1 = pd.read_csv(group1.doc_logfc.path,  index_col='SYMBOL', sep='\t')
                    df_2 = pd.read_csv(group2.doc_logfc.path,  index_col='SYMBOL', sep='\t') 
                    df_3 = pd.read_csv(group3.doc_logfc.path,  index_col='SYMBOL', sep='\t') 
                    
                if (df_1['adj.P.Val']>pval_tres).all():
                    df_1 = df_1[(np.absolute(df_1['logFC'])>logFC_tres)]
                else:               
                    df_1 = df_1[(df_1['adj.P.Val']<pval_tres) & (np.absolute(df_1['logFC'])>logFC_tres)]
                if (df_2['adj.P.Val']>pval_tres).all():
                    df_2 = df_2[(np.absolute(df_2['logFC'])>logFC_tres)]
                else:
                    df_2 = df_2[(df_2['adj.P.Val']<pval_tres) & (np.absolute(df_2['logFC'])>logFC_tres)]
                if (df_3['adj.P.Val']>pval_tres).all():
                    df_3 = df_3[(np.absolute(df_3['logFC'])>logFC_tres)]
                else:
                    df_3 = df_3[(df_3['adj.P.Val']<pval_tres) & (np.absolute(df_3['logFC'])>logFC_tres)]                
                
                df_1 = pd.DataFrame(df_1['logFC'])
                df_2 = pd.DataFrame(df_2['logFC'])
                df_3 = pd.DataFrame(df_3['logFC'])
                df_1.columns = ['0']
                df_2.columns = ['0']
                df_3.columns = ['0']
            
            elif path_gene == 'deeplearning': # For GSK prj2 only! not dynamic!
            #raise Exception('extended')
                if is_metabolic=='true':
                    df_1 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/New_top_model_FC_whole.csv',
                                  index_col='probeset')
                    df_2 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/New_top_model_FC_EC.csv',
                                  index_col='probeset')
                    df_3 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/New_top_model_FC_SM.csv',
                                  index_col='probeset')
                else:
                    df_1 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/New_top_model_FC_ext_whole.csv',
                                  index_col='probeset')
                    df_2 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/New_top_model_FC_ext_EC.csv',
                                  index_col='probeset')
                    df_3 = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/Extended/New_top_model_FC_ext_SM.csv',
                                  index_col='probeset')
                    
                    
                df_1 = df_1.sort(columns='mean', ascending=False).head(1000)
                df_2 = df_2.sort(columns='mean', ascending=False).head(1000)
                df_3 = df_3.sort(columns='mean', ascending=False).head(1000)
            
                df_1.rename(columns={'mean': '0'}, inplace=True)
                df_2.rename(columns={'mean': '0'}, inplace=True)
                df_3.rename(columns={'mean': '0'}, inplace=True)
                
                #raise Exception('deeplearning')
            
            s_tumour1 = df_1['0']#.round(decimals=2)
            s_tumour1.name = lMembers[0]
           
            s_tumour2 = df_2['0']#.round(decimals=2)
            s_tumour2.name = lMembers[1]
            
            s_tumour3 = df_3['0']#.round(decimals=2)
            s_tumour3.name = lMembers[2]
            
            if regulation == 'updown':
                s_tumour1_up = s_tumour1[s_tumour1>0]
                s_tumour1_down = s_tumour1[s_tumour1<0]
                s_tumour2_up = s_tumour2[s_tumour2>0]
                s_tumour2_down = s_tumour2[s_tumour2<0]
                s_tumour3_up = s_tumour3[s_tumour3>0]
                s_tumour3_down = s_tumour3[s_tumour3<0]
                
                df_up = pd.DataFrame(s_tumour1_up).join(pd.DataFrame(s_tumour2_up), how='inner', sort=True)
                df_up = df_up.join(pd.DataFrame(s_tumour3_up), how='inner', sort=True)         
                df_down = pd.DataFrame(s_tumour1_down).join(pd.DataFrame(s_tumour2_down), how='inner', sort=True)
                df_down = df_down.join(pd.DataFrame(s_tumour3_down), how='inner', sort=True) 
                joined_df = df_up.append(df_down)
            elif regulation == 'up':            
                s_tumour1 = s_tumour1[s_tumour1>0]
                s_tumour2 = s_tumour2[s_tumour2>0]
                s_tumour3 = s_tumour3[s_tumour3>0]
                joined_df = pd.DataFrame(s_tumour1).join(pd.DataFrame(s_tumour2), how='inner')
                joined_df = joined_df.join(pd.DataFrame(s_tumour3), how='inner')
                
                if path_gene =='deeplearning':
                    joined_df = joined_df.join(df_1['genes'], how='inner')
                    joined_df.reset_index(inplace=True)
                    joined_df.set_index('genes', inplace=True)
                    #raise Exception('deeplearning')
            elif regulation == 'down':  
                s_tumour1 = s_tumour1[s_tumour1<0]
                s_tumour2 = s_tumour2[s_tumour2<0]
                s_tumour3 = s_tumour3[s_tumour3<0]
                joined_df = pd.DataFrame(s_tumour1).join(pd.DataFrame(s_tumour2), how='inner')
                joined_df = joined_df.join(pd.DataFrame(s_tumour3), how='inner')
            
            if path_gene !='deeplearning':
                joined_df = joined_df.groupby(joined_df.index, level=0).mean()
            
            joined_df.reset_index(inplace=True)
            df_json = joined_df.to_json(orient='values')
        
        
        elif inter_num == 4:
            if path_gene == 'pathways':
                group1 = PathwayGroup.objects.get(name=lMembers[0], report=report)
                df_1 = pd.read_csv(group1.doc_proc.path, index_col='Pathway')
                df_1 = df_1[np.absolute(df_1['0'])>pas_tres]
                
                group2 = PathwayGroup.objects.get(name=lMembers[1], report=report)
                df_2 = pd.read_csv(group2.doc_proc.path, index_col='Pathway')
                df_2 = df_2[np.absolute(df_2['0'])>pas_tres]
                
                group3 = PathwayGroup.objects.get(name=lMembers[2], report=report)
                df_3 = pd.read_csv(group3.doc_proc.path,  index_col='Pathway')
                df_3 = df_3[np.absolute(df_3['0'])>pas_tres]
                
                group4 = PathwayGroup.objects.get(name=lMembers[3], report=report)
                df_4 = pd.read_csv(group4.doc_proc.path,  index_col='Pathway')
                df_4 = df_4[np.absolute(df_4['0'])>pas_tres]
                
                if is_metabolic=='true':
                    df_1 = df_1[df_1['Database']=='metabolism']
                    df_2 = df_2[df_2['Database']=='metabolism']
                    df_3 = df_3[df_3['Database']=='metabolism']
                    df_4 = df_4[df_4['Database']=='metabolism']
                else:                
                    df_1 = df_1[df_1['Database']!='metabolism']
                    df_2 = df_2[df_2['Database']!='metabolism']
                    df_3 = df_3[df_3['Database']!='metabolism']
                    df_4 = df_4[df_4['Database']!='metabolism']
            
            elif path_gene =='genes':
                group1 = GeneGroup.objects.get(name=lMembers[0], report=report)
                group2 = GeneGroup.objects.get(name=lMembers[1], report=report)
                group3 = GeneGroup.objects.get(name=lMembers[2], report=report)
                group4 = GeneGroup.objects.get(name=lMembers[3], report=report)
                
                try:
                    df_1 = pd.read_csv(group1.doc_logfc.path,  index_col='SYMBOL')
                    df_2 = pd.read_csv(group2.doc_logfc.path,  index_col='SYMBOL') 
                    df_3 = pd.read_csv(group3.doc_logfc.path,  index_col='SYMBOL')
                    df_4 = pd.read_csv(group4.doc_logfc.path,  index_col='SYMBOL')
                except:
                    df_1 = pd.read_csv(group1.doc_logfc.path,  index_col='SYMBOL', sep='\t')
                    df_2 = pd.read_csv(group2.doc_logfc.path,  index_col='SYMBOL', sep='\t') 
                    df_3 = pd.read_csv(group3.doc_logfc.path,  index_col='SYMBOL', sep='\t')
                    df_4 = pd.read_csv(group4.doc_logfc.path,  index_col='SYMBOL', sep='\t')
                
                df_1.replace([np.inf, -np.inf], 0, inplace=True)                
                df_2.replace([np.inf, -np.inf], 0, inplace=True)                
                df_3.replace([np.inf, -np.inf], 0, inplace=True)                 
                df_4.replace([np.inf, -np.inf], 0, inplace=True)
                
                if (df_1['adj.P.Val']>pval_tres).all():
                    df_1 = df_1[(np.absolute(df_1['logFC'])>logFC_tres)]
                else:               
                    df_1 = df_1[(df_1['adj.P.Val']<pval_tres) & (np.absolute(df_1['logFC'])>logFC_tres)]
                if (df_2['adj.P.Val']>pval_tres).all():
                    df_2 = df_2[(np.absolute(df_2['logFC'])>logFC_tres)]
                else:
                    df_2 = df_2[(df_2['adj.P.Val']<pval_tres) & (np.absolute(df_2['logFC'])>logFC_tres)]
                if (df_3['adj.P.Val']>pval_tres).all():
                    df_3 = df_3[(np.absolute(df_3['logFC'])>logFC_tres)]
                else:
                    df_3 = df_3[(df_3['adj.P.Val']<pval_tres) & (np.absolute(df_3['logFC'])>logFC_tres)]    
                if (df_4['adj.P.Val']>pval_tres).all():
                    df_4 = df_4[(np.absolute(df_4['logFC'])>logFC_tres)]
                else:
                    df_4 = df_4[(df_4['adj.P.Val']<pval_tres) & (np.absolute(df_4['logFC'])>logFC_tres)]                 
                
                df_1 = pd.DataFrame(df_1['logFC'])
                df_2 = pd.DataFrame(df_2['logFC'])
                df_3 = pd.DataFrame(df_3['logFC'])
                df_4 = pd.DataFrame(df_4['logFC'])
                df_1.columns = ['0']
                df_2.columns = ['0']
                df_3.columns = ['0']
                df_4.columns = ['0']
            
            
            s_tumour1 = df_1['0'].round(decimals=2)
            s_tumour1.name = lMembers[0]
           
            s_tumour2 = df_2['0'].round(decimals=2)
            s_tumour2.name = lMembers[1]
            
            s_tumour3 = df_3['0'].round(decimals=2)
            s_tumour3.name = lMembers[2]
            
            s_tumour4 = df_4['0'].round(decimals=2)
            s_tumour4.name = lMembers[3]
            
            if regulation == 'updown':
                s_tumour1_up = s_tumour1[s_tumour1>0]
                s_tumour1_down = s_tumour1[s_tumour1<0]
                s_tumour2_up = s_tumour2[s_tumour2>0]
                s_tumour2_down = s_tumour2[s_tumour2<0]
                s_tumour3_up = s_tumour3[s_tumour3>0]
                s_tumour3_down = s_tumour3[s_tumour3<0]
                s_tumour4_up = s_tumour4[s_tumour4>0]
                s_tumour4_down = s_tumour4[s_tumour4<0]
                
                df_up = pd.DataFrame(s_tumour1_up).join(pd.DataFrame(s_tumour2_up), how='inner', sort=True)
                df_up = df_up.join(pd.DataFrame(s_tumour3_up), how='inner', sort=True)
                df_up = df_up.join(pd.DataFrame(s_tumour4_up), how='inner', sort=True)         
                df_down = pd.DataFrame(s_tumour1_down).join(pd.DataFrame(s_tumour2_down), how='inner', sort=True)
                df_down = df_down.join(pd.DataFrame(s_tumour3_down), how='inner', sort=True)
                df_down = df_down.join(pd.DataFrame(s_tumour4_down), how='inner', sort=True) 
                joined_df = df_up.append(df_down)
            elif regulation == 'up':            
                s_tumour1 = s_tumour1[s_tumour1>0]
                s_tumour2 = s_tumour2[s_tumour2>0]
                s_tumour3 = s_tumour3[s_tumour3>0]
                s_tumour4 = s_tumour4[s_tumour4>0]
                joined_df = pd.DataFrame(s_tumour1).join(pd.DataFrame(s_tumour2), how='inner')
                joined_df = joined_df.join(pd.DataFrame(s_tumour3), how='inner')
                joined_df = joined_df.join(pd.DataFrame(s_tumour4), how='inner')
            elif regulation == 'down':  
                s_tumour1 = s_tumour1[s_tumour1<0]
                s_tumour2 = s_tumour2[s_tumour2<0]
                s_tumour3 = s_tumour3[s_tumour3<0]
                s_tumour4 = s_tumour4[s_tumour4<0]
                joined_df = pd.DataFrame(s_tumour1).join(pd.DataFrame(s_tumour2), how='inner')
                joined_df = joined_df.join(pd.DataFrame(s_tumour3), how='inner')
                joined_df = joined_df.join(pd.DataFrame(s_tumour4), how='inner')
            
            joined_df = joined_df.groupby(joined_df.index, level=0).mean()
            joined_df.reset_index(inplace=True)
            df_json = joined_df.to_json(orient='values')
        #raise Exception('venn table')       
        response_data = {'aaData': json.loads(df_json)}
        return HttpResponse(json.dumps(response_data), content_type="application/json")


class ReportAjaxPathwayVennTableGSK(TemplateView):
    """
    Used for venn table in GSK project 2 only
    """
    template_name="website/report_ajax_venn.html"
    
    
    def dispatch(self, request, *args, **kwargs):
        return super(ReportAjaxPathwayVennTableGSK, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):        
        
        report = Report.objects.get(pk=self.request.GET.get('reportID'))
        group = self.request.GET.get('group')
        
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/predictors_PAS_NN.csv')
        #raise Exception('stop')
        #ec up
        if group == 'toppath_ec_w_up':
            df = df[df['comparison'] == 'Endothelial Cells' ]
            df = df[df['group'] == 'pathways_only' ]
            df = df[df['direction'] == 1 ]
        if group == 'topprobe_ec_w_up':
            df = df[df['comparison'] == 'Endothelial Cells' ]
            df = df[df['group'] == 'NN_only' ]
            df = df[df['direction'] == 1 ]
        if group == 'toppath_topprobe_ec_w_up':
            df = df[df['comparison'] == 'Endothelial Cells' ]
            df = df[df['group'] == 'both' ]
            df = df[df['direction'] == 1 ]
        #ec down
        if group == 'toppath_ec_w_down':
            df = df[df['comparison'] == 'Endothelial Cells' ]
            df = df[df['group'] == 'pathways_only' ]
            df = df[df['direction'] == -1 ]
        if group == 'topprobe_ec_w_down':
            df = df[df['comparison'] == 'Endothelial Cells' ]
            df = df[df['group'] == 'NN_only' ]
            df = df[df['direction'] == -1 ]
        if group == 'toppath_topprobe_ec_w_down':
            df = df[df['comparison'] == 'Endothelial Cells' ]
            df = df[df['group'] == 'both' ]
            df = df[df['direction'] == -1 ]    
        
        #sm up    
        if group == 'toppath_sm_w_up':
            df = df[df['comparison'] == 'Smooth Muscle' ]
            df = df[df['group'] == 'pathways_only' ]
            df = df[df['direction'] == 1 ]
        if group == 'topprobe_sm_w_up':
            df = df[df['comparison'] == 'Smooth Muscle' ]
            df = df[df['group'] == 'NN_only' ]
            df = df[df['direction'] == 1 ]
        if group == 'toppath_topprobe_sm_w_up':
            df = df[df['comparison'] == 'Smooth Muscle' ]
            df = df[df['group'] == 'both' ]
            df = df[df['direction'] == 1 ]
        
        #sm down    
        if group == 'toppath_sm_w_down':
            df = df[df['comparison'] == 'Smooth Muscle' ]
            df = df[df['group'] == 'pathways_only' ]
            df = df[df['direction'] == -1 ]
        if group == 'topprobe_sm_w_down':
            df = df[df['comparison'] == 'Smooth Muscle' ]
            df = df[df['group'] == 'NN_only' ]
            df = df[df['direction'] == -1 ]
        if group == 'toppath_topprobe_sm_w_down':
            df = df[df['comparison'] == 'Smooth Muscle' ]
            df = df[df['group'] == 'both' ]
            df = df[df['direction'] == -1 ]
            
        # EXTENDED
        #ec up
        if group == 'toppath_ec_e_up':
            df = df[df['comparison'] == 'Endothelial Cells (ext)' ]
            df = df[df['group'] == 'pathways_only' ]
            df = df[df['direction'] == 1 ]
        if group == 'topprobe_ec_e_up':
            df = df[df['comparison'] == 'Endothelial Cells (ext)' ]
            df = df[df['group'] == 'NN_only' ]
            df = df[df['direction'] == 1 ]
        if group == 'toppath_topprobe_ec_e_up':
            df = df[df['comparison'] == 'Endothelial Cells (ext)' ]
            df = df[df['group'] == 'both' ]
            df = df[df['direction'] == 1 ]
        #ec down
        if group == 'toppath_ec_e_down':
            df = df[df['comparison'] == 'Endothelial Cells (ext)' ]
            df = df[df['group'] == 'pathways_only' ]
            df = df[df['direction'] == -1 ]
        if group == 'topprobe_ec_e_down':
            df = df[df['comparison'] == 'Endothelial Cells (ext)' ]
            df = df[df['group'] == 'NN_only' ]
            df = df[df['direction'] == -1 ]
        if group == 'toppath_topprobe_ec_e_down':
            df = df[df['comparison'] == 'Endothelial Cells (ext)' ]
            df = df[df['group'] == 'both' ]
            df = df[df['direction'] == -1 ]
            
        #sm up
        if group == 'toppath_sm_e_up':
            df = df[df['comparison'] == 'Smooth Muscle (ext)' ]
            df = df[df['group'] == 'pathways_only' ]
            df = df[df['direction'] == 1 ]
        if group == 'topprobe_sm_e_up':
            df = df[df['comparison'] == 'Smooth Muscle (ext)' ]
            df = df[df['group'] == 'NN_only' ]
            df = df[df['direction'] == 1 ]
        if group == 'toppath_topprobe_sm_e_up':
            df = df[df['comparison'] == 'Smooth Muscle (ext)' ]
            df = df[df['group'] == 'both' ]
            df = df[df['direction'] == 1 ]
        #ec down
        if group == 'toppath_sm_e_down':
            df = df[df['comparison'] == 'Smooth Muscle (ext)' ]
            df = df[df['group'] == 'pathways_only' ]
            df = df[df['direction'] == -1 ]
        if group == 'topprobe_sm_e_down':
            df = df[df['comparison'] == 'Smooth Muscle (ext)' ]
            df = df[df['group'] == 'NN_only' ]
            df = df[df['direction'] == -1 ]
        if group == 'toppath_topprobe_sm_e_down':
            df = df[df['comparison'] == 'Smooth Muscle (ext)' ]
            df = df[df['group'] == 'both' ]
            df = df[df['direction'] == -1 ]     
        
        
        
        
        df_json = df[['comparison', 'gene']].to_json(orient='values')
        #raise Exception('venn table')       
        response_data = {'data': json.loads(df_json)}
        return HttpResponse(json.dumps(response_data), content_type="application/json")

import matplotlib.pyplot as plt
import matplotlib.colors as colors
import matplotlib.cm as cmx
import matplotlib as mpl
import struct

def shiftedColorMap(cmap, start=0, midpoint=0, stop=1.0, name='shiftedcmap'):

    cdict = {
        'red': [],
        'green': [],
        'blue': [],
        'alpha': []
    }

    # regular index to compute the colors
    reg_index = np.linspace(start, stop, 257)

    # shifted index to match the data
    shift_index = np.hstack([
        np.linspace(0.0, midpoint, 128, endpoint=False), 
        np.linspace(midpoint, 1.0, 129, endpoint=True)
    ])

    for ri, si in zip(reg_index, shift_index):
        r, g, b, a = cmap(ri)

        cdict['red'].append((si, r, r))
        cdict['green'].append((si, g, g))
        cdict['blue'].append((si, b, b))
        cdict['alpha'].append((si, a, a))

    newcmap = mpl.colors.LinearSegmentedColormap(name, cdict)
    plt.register_cmap(cmap=newcmap)

    return newcmap
    
class ReportAjaxPathDetail(TemplateView):
    template_name = 'report/report_ajax_pathway_detail.html'
    
   
    
    def dispatch(self, request, *args, **kwargs):
        return super(ReportAjaxPathDetail, self).dispatch(request, *args, **kwargs)
    
    def get_context_data(self, *args, **kwargs):
        context = super(ReportAjaxPathDetail, self).get_context_data(**kwargs)
        
        report = Report.objects.get(pk=self.request.GET['reportID'])
        path_name = self.request.GET['pathway']
        group_name = self.request.GET['group_name']
        
        group_name = group_name.replace(' p-value', '')
        
        try:
            group = GeneGroup.objects.get(report=report, name=group_name)
        except:
            pass
        organism = self.request.GET['organism']
        
        pval_tres = report.pval_theshold_compare
        logFC_tres = report.logcf_theshold_compare
        pas_tres = report.pas_theshold       

        
        db_to_exclude = ['primary_old', 'kegg', 'kegg_10', 'kegg_adjusted_10']
        
        if report.id==2:
            pathway = Pathway.objects.filter(organism=organism, name=path_name).exclude(database__in=['primary_old'])[0]
        else: 
            try:
                pathway = Pathway.objects.filter(organism=organism, name=path_name).exclude(database__in=db_to_exclude)[0]
            except:
                try:
                    name=str(self.request.GET['pathway'].replace('.', '-'))
                    
                    pathway = Pathway.objects.filter(organism=organism, name=name).exclude(database__in=db_to_exclude)[0]
                except:
                    name=str(self.request.GET['pathway'].replace('.', '('))
                    name = name[:-1] + ')'
                    pathway = Pathway.objects.filter(organism=organism, name=name).exclude(database__in=db_to_exclude)[0]
                      
                    
                
        gene_data = []
        for gene in pathway.gene_set.all():
                nodes = ','.join([str(i) for i in Node.objects.filter(pathway=pathway, component__name=gene.name).distinct()])
                
                gene_data.append({'SYMBOL':gene.name.strip().upper(),
                                  'Node(s)':nodes })   
           
        gene_df = pd.DataFrame(gene_data).set_index('SYMBOL')
            
        try:
            df_file_cnr = pd.read_csv(group.doc_logfc.path, index_col='SYMBOL')
            df_file_cnr.rename(columns={'log2FoldChange': 'logFC', 'padj': 'adj.P.Val'}, inplace=True)
        except:
            try:
                df_file_cnr = pd.read_csv(group.doc_logfc.path, index_col='SYMBOL', sep='\t')
                df_file_cnr.rename(columns={'log2FoldChange': 'logFC', 'padj': 'adj.P.Val'}, inplace=True)
            except: # for Henkel reports only!!!
                df_file_cnr = pd.read_csv(settings.MEDIA_ROOT+"/report-portal/"+report.slug+"/Henkel_mean.csv", 
                                      sep=None, index_col='SYMBOL')
                df_file_cnr = df_file_cnr[['logFC', 'adj.P.Val']]
        
        if (df_file_cnr['adj.P.Val']>pval_tres).all():
                    df_file_cnr = df_file_cnr[(np.absolute(df_file_cnr['logFC'])>logFC_tres)]
                    
                    
        else:
            if report.id == 7:   # for Above & Beyond only     
                df_file_cnr = df_file_cnr[(df_file_cnr['pvalue']<pval_tres) & (np.absolute(df_file_cnr['logFC'])>logFC_tres)] #take all the genes
            else:
                df_file_cnr = df_file_cnr[(df_file_cnr['adj.P.Val']<pval_tres) & (np.absolute(df_file_cnr['logFC'])>logFC_tres)]
        
        
        df_file_cnr = df_file_cnr['logFC'].round(decimals=2)
       
        
        df_file_cnr.name = 'log2(Fold-change)' #log2(Fold-change)
        
        joined_df = gene_df.join(df_file_cnr, how='inner')
        joined_df.reset_index(inplace=True)
        
        joined_df.index += 1
        
        joined_df.fillna(0, inplace=True)
        
        context['joined'] = joined_df[['SYMBOL', 'Node(s)', 'log2(Fold-change)']].to_html(classes=['table', 
                                                                                                   'table-bordered', 
                                                                                                   'table-striped', 
                                                                                                   'genes_for_path'])
        context['diff_genes_count'] = len(joined_df.index)
         
        nComp = []
        
        G=nx.DiGraph()
        
        for _, row in joined_df.iterrows():                
            
            loComp = Component.objects.filter(name = row['SYMBOL'], node__pathway=pathway)
            
            for comp in loComp:
                comp.cnr = np.power(2, row['log2(Fold-change)'])
                nComp.append(comp)
                    
        lNodes = []
        lNEL = []        
        
        for node in pathway.node_set.all():
            
            lcurrentComponents = [i.name for i in node.component_set.all()]
            
            node.nel = 0.0
            node.numDiffComp = 0
            node.sumDiffComp = 0.0
            node.color = "grey"
            node.strokeWidth = 1
            for component in nComp:
                if component.name in lcurrentComponents: 
                    if component.cnr != 0:
                        node.numDiffComp += 1
                        node.sumDiffComp += float(component.cnr)
            if node.numDiffComp >0 :
                node.nel = node.sumDiffComp / node.numDiffComp
                lNEL.append(node.sumDiffComp / node.numDiffComp)
            lNodes.append(node)
                
        if lNEL: #check if list is not empty
            #choosing colormap for static image
            lNEL = np.log(lNEL)
            mmin = np.min(lNEL)
            mmax = np.max(np.absolute(lNEL)) # absolute was made for Aliper special, remove if needed
            mid = 1 - mmax/(mmax + abs(mmin))
        
            if mmax<0 and mmin<0:                    
                shifted_cmap = plt.get_cmap('Reds_r')
                mmax = 0
            if  mmax>0 and mmin>0:
                shifted_cmap = plt.get_cmap('Greens')
                mmin = 0
            else:  
                cmap = plt.get_cmap('PiYG')
                shifted_cmap = shiftedColorMap(cmap, start=0, midpoint=mid, stop=1, name='shrunk')
            cNormp  = colors.Normalize(vmin=mmin, vmax=mmax)
            scalarMap = cmx.ScalarMappable(norm=cNormp, cmap=shifted_cmap)
        
        finalNodes = []
        for nod in lNodes:
            if nod.nel > 1:
                nod.color = "green"
                nod.strokeWidth = np.log2(nod.nel)
            if nod.nel <= 1 and nod.nel > 0:
                nod.color = "red"
                nod.strokeWidth = np.log2(nod.nel)
            finalNodes.append(nod)
            
            if nod.nel!=0:
                
                ffil = "#"+struct.pack('BBB',*scalarMap.to_rgba(np.log(nod.nel), bytes=True)[:3]).encode('hex').upper()
            else:
                ffil = "grey"
            G.add_node(nod.name, color='black',style='filled',
                               fillcolor=ffil)            
        
        context['colorNodes'] = finalNodes              
        
        dRelations = []
        for node in pathway.node_set.all():
            for inrel in node.inrelations.all():
                relColor = 'black'
                if inrel.reltype == '1':
                    relColor = 'green'
                if inrel.reltype == '0':
                    relColor = 'red'
                dRelations.append({ inrel.fromnode.name : [inrel.tonode.name, relColor] })
                G.add_edge(inrel.fromnode.name.encode('ascii','ignore'), inrel.tonode.name.encode('ascii','ignore'), color=relColor)      
        
        # DRAW STATIC IMAGE
        A=nx.to_agraph(G)
        A.layout(prog='dot')
        A.draw(settings.MEDIA_ROOT+"/pathway.svg")
        
        context['pathway'] = pathway
        context['dRelations'] = dRelations
        
        context['rand'] = random.random() 
        
        context['report'] = report
        context['categories'] = group_name
        
        #raise Exception('path details')
        
        return context    


class ReportTfTableJson(TemplateView): 
    template_name="report/report_detail.html"
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(ReportTfTableJson, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        
        file_name = request.GET.get('file_name')            
        
        try:    
            df_tf = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name,
                                  sep=' ')
        except:
            df_tf = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name, sep=None
                                  )
                
        df_tf.fillna(1, inplace=True)                 
         
        df_tf = df_tf[(df_tf['padj']<0.05)]         
        
        df_tf['padj'] = df_tf['padj'].map('{:,.2e}'.format)
        df_tf['p.value'] = df_tf['p.value'].map('{:,.2e}'.format)
        
        df_tf['padj'] = df_tf['padj'].apply(str)
        df_tf['p.value'] = df_tf['p.value'].apply(str)
        
        output_json = df_tf.to_json(orient='values')
        #raise Exception('tf') 
        response_data = {'data': json.loads(output_json)}
        
        return HttpResponse(json.dumps(response_data), content_type="application/json") 

class ReportTfTrrustTableJson(TemplateView): 
    template_name="report/report_detail.html"
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(ReportTfTrrustTableJson, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        
        file_name = request.GET.get('file_name')            
        
        try:    
            df_tf = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name,
                                  sep=' ')
        except:
            df_tf = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name,
                                  )
                
        df_tf.fillna(1, inplace=True)                 
         
        df_tf = df_tf[(df_tf['padj']<0.05)]         
        
        df_tf['padj'] = df_tf['padj'].map('{:,.2e}'.format)
        df_tf['p.value'] = df_tf['p.value'].map('{:,.2e}'.format)
        
        df_tf['padj'] = df_tf['padj'].apply(str)
        df_tf['p.value'] = df_tf['p.value'].apply(str)
        
        output_json = df_tf.to_json(orient='values')
        #raise Exception('tf') 
        response_data = {'data': json.loads(output_json)}
        
        return HttpResponse(json.dumps(response_data), content_type="application/json") 

class ReportAjaxTfDetail(TemplateView):
    template_name = 'report/report_ajax_tf_detail.html'
    
   
    
    def dispatch(self, request, *args, **kwargs):
        return super(ReportAjaxTfDetail, self).dispatch(request, *args, **kwargs)
    
    def get_context_data(self, *args, **kwargs):
        context = super(ReportAjaxTfDetail, self).get_context_data(**kwargs)
        
        report = Report.objects.get(pk=self.request.GET['reportID'])
        group = TfGroup.objects.get(report=report, name=self.request.GET['group_name'])
        tf_name = self.request.GET['tf_name']
        
        df_tf = pd.read_csv(group.document.path, sep=' ')
        df_tf.fillna(1, inplace=True)      
        df_tf = df_tf[(df_tf['padj']<0.05)]
        ldifferential = df_tf['factor'].tolist()
        
        
        # Get logFC values to color nodes
        gene_group = GeneGroup.objects.get(report=report, name=self.request.GET['group_name'])
        try:
            df_logfc = pd.read_csv(gene_group.doc_logfc.path, sep=None)                        
            df_logfc = df_logfc[['SYMBOL', 'logFC']]
        except:
            df_logfc = pd.read_csv(gene_group.doc_logfc.path, sep=None)                        
            df_logfc = df_logfc[['SYMBOL', 'log2FoldChange', 'padj']]
            df_logfc.rename(columns={'log2FoldChange': 'logFC', 'padj': 'adj.P.Val'}, inplace=True)
            
        df_logfc = df_logfc[df_logfc['SYMBOL'].isin(ldifferential)]
        df_logfc.dropna(inplace=True)
        df_logfc.set_index('SYMBOL', inplace=True)
        df_logfc.index.name = 'to'
        
        
        
        ##### Use specific file for each cell line
        if 'A549' in group.name:
            specific_df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/adenocarcinoma_cell_line.txt',
                                    sep='\t' , header=None)
            s_file_name = 'adenocarcinoma_cell_line'
        elif 'MCF7' in group.name:
            specific_df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/breast_carcinoma_cell_line.txt',
                                    sep='\t' , header=None)
            s_file_name = 'breast_carcinoma_cell_line'
        
        specific_df.columns = ['from', 'to', 'source']
        specific_df = specific_df[specific_df['from']==tf_name]
        
        lStargets = specific_df['to'].tolist()  
        s_intersection =  list(set(lStargets).intersection(ldifferential))
        specific_intersect = specific_df[specific_df['to'].isin(s_intersection)]
        
        # join with logfc
        specific_intersect.set_index('to', inplace=True)
        
        specific_joined = specific_intersect.join(df_logfc, how='inner')
        specific_joined.reset_index(inplace=True)
        
        lNEL = specific_joined['logFC']
        mmin = np.min(lNEL)
        mmax = np.max(np.absolute(lNEL)) # absolute was made for Aliper special, remove if needed
        mid = 1 - mmax/(mmax + abs(mmin))
        
        if mmax<0 and mmin<0:                    
            shifted_cmap = plt.get_cmap('Reds_r')
            mmax = 0
        if  mmax>0 and mmin>0:
            shifted_cmap = plt.get_cmap('Greens')
            mmin = 0
        else:  
            cmap = plt.get_cmap('PiYG')
            shifted_cmap = shiftedColorMap(cmap, start=0, midpoint=mid, stop=1, name='shrunk')
        cNormp  = colors.Normalize(vmin=mmin, vmax=mmax)
        scalarMap = cmx.ScalarMappable(norm=cNormp, cmap=shifted_cmap)
        
        #raise Exception('TF detail')
        # Draw Graph
        
        G_spe=nx.DiGraph()
        
        for index, row in specific_joined.iterrows():
            
            if row['logFC']!=0:                
                ffil = "#"+struct.pack('BBB',*scalarMap.to_rgba(row['logFC'], bytes=True)[:3]).encode('hex').upper()
            else:
                ffil = "grey"
            
            G_spe.add_node(row['to'], color='black',style='filled',
                               fillcolor=ffil)
            
            relColor = 'black'
            """  As there is no relation type in specific file!!!!
            if row['relType'] == 'Activation':
                relColor = 'green'
            if row['relType'] == 'Repression':
                relColor = 'red'
            """
            G_spe.add_edge(row['from'], row['to'], color=relColor)
            
        
        # DRAW STATIC IMAGE
        A_spe=nx.to_agraph(G_spe)
        A_spe.layout(prog='circo')
        A_spe.draw(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/tf_specific.svg')
        
        #raise Exception('TF detail')
        
        ##### Use common file for each cell line
        common_df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/trrust_rawdata.txt',
                                  sep='\t', header=None)        
        common_df.columns = ['from', 'to', 'relType', 'source']
        
        common_df = common_df[common_df['from']==tf_name] #filter related to current tf
        
        lCtargets = common_df['to'].tolist()  
        c_intersection =  list(set(lCtargets).intersection(ldifferential))
        common_intersect = common_df[common_df['to'].isin(c_intersection)]
        
        # join with logfc
        common_intersect.set_index('to', inplace=True)
        
        common_joined = common_intersect.join(df_logfc, how='inner')
        common_joined.reset_index(inplace=True)
        
        
        lNEL = common_joined['logFC']
        mmin = np.min(lNEL)
        mmax = np.max(np.absolute(lNEL)) # absolute was made for Aliper special, remove if needed
        mid = 1 - mmax/(mmax + abs(mmin))
        
        if mmax<0 and mmin<0:                    
            shifted_cmap = plt.get_cmap('Reds_r')
            mmax = 0
        if  mmax>0 and mmin>0:
            shifted_cmap = plt.get_cmap('Greens')
            mmin = 0
        else:  
            cmap = plt.get_cmap('PiYG')
            shifted_cmap = shiftedColorMap(cmap, start=0, midpoint=mid, stop=1, name='shrunk')
        cNormp  = colors.Normalize(vmin=mmin, vmax=mmax)
        scalarMap = cmx.ScalarMappable(norm=cNormp, cmap=shifted_cmap)
       
        # Draw Graph
        
        G_com=nx.DiGraph()
        
        for index, row in common_joined.iterrows():
            
            if row['logFC']!=0:                
                ffil = "#"+struct.pack('BBB',*scalarMap.to_rgba(row['logFC'], bytes=True)[:3]).encode('hex').upper()
            else:
                ffil = "grey"
            
            G_com.add_node(row['to'], color='black',style='filled',
                               fillcolor=ffil)
                       
            relColor = 'black'
            if row['relType'] == 'Activation':
                relColor = 'green'
            if row['relType'] == 'Repression':
                relColor = 'red'
            
            G_com.add_edge(row['from'], row['to'], color=relColor)
            
        
        # DRAW STATIC IMAGE
        A_com=nx.to_agraph(G_com)
        A_com.layout(prog='dot')
        A_com.draw(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/tf_common.svg')
        
        context['s_file_name'] = s_file_name
        context['c_file_name'] = 'trrust_rawdata'
        context['report_slug'] = report.slug
        context['tf_name'] = tf_name
        context['rand'] = random.random()
        
        #raise Exception('TF detail')
        
        return context

   
        
class ReportDlFarmJson(TemplateView):
    template_name="report/report_detail.html"
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(ReportDlFarmJson, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        
        report = Report.objects.get(pk=self.request.GET['report_id'])
        group = self.request.GET['group_name']
        
        file_type = self.request.GET['type']
        
        
        deeplearning = report.deeplearning_set.all()[0]
        
        if file_type == 'fc':
            df_farm = pd.read_csv(deeplearning.farmclass.path, index_col='Name', sep='\t')
            s_threshold = df_farm.iloc[0, 7:]
        elif file_type == 'se':
            df_farm = pd.read_csv(deeplearning.sideeff.path, index_col='Name', sep='\t')
            s_threshold = df_farm.iloc[0, 33:]
        
        lGname = group.split('_')
        
        if lGname[0]!='overall':
            lGname[0] = lGname[0].replace('D', '') # in case of groups not overall 
        
        
        if lGname[0]!='overall':
            df_farm = df_farm[(df_farm['compound']== int(lGname[0])) & 
                           (df_farm['concentration']==int(lGname[1])) &
                            (df_farm['celltype']==lGname[2] )]
        else:
            df_farm = df_farm[(df_farm['compound']== int(lGname[1]))]
        
       
        
        
        
        if file_type == 'fc':
            s_features = df_farm.iloc[:, 7:] # leave only columns with features, excluding 'Best threshold'
        if file_type == 'se':
            s_features = df_farm.iloc[:, 33:] # leave only columns with features, excluding 'Best threshold'
        
        
        s_features = s_features.mean()        
        s_features.sort(ascending=False)
        
        """
        aaa = s_features['Antineoplastic Agents']
        
        bbb = s_features[s_features>0.5].count()
        
        raise Exception('stop')
        """
        
        if file_type == 'se': # leave only top-30 Side Effects
            s_features = s_features[:30]
        
        
        barplot_data = []
        threshold = []
        for index, value in s_features.iteritems():
            value = "%.2f" % value
            if float(value)>0:
                barplot_data.append([index, float(value)])
                threshold.append(s_threshold[index])
            
        
        
        
        #raise Exception('stop11')
        if file_type == 'fc':
            plot_title = 'Therapeutic use'
            
        elif file_type == 'se':
            plot_title = 'Side Effects'
            
            
        if lGname[0]!='overall':
            plot_subtitle = group
        else:
            plot_subtitle = plot_subtitle = 'Overall for Drug ' + lGname[1]     
        
        output = {'barplot': barplot_data,
                  'threshold': threshold,
                  'title': plot_title,
                  'subtitle': plot_subtitle
                  }
        
        response_data =  json.dumps(output)
        return HttpResponse(json.dumps(response_data), content_type="application/json") 
    
    
class ReportSimilarityJson(TemplateView): 
    template_name="report/report_detail.html"
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(ReportSimilarityJson, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        report = Report.objects.get(pk=self.request.GET['reportID'])        
        
        similarity = report.deeplearning_set.all()[1]
        sim_group_name = self.request.GET['sim_group_name']            
        
        lsim_group = sim_group_name.split('_')
        lsim_group[0] = lsim_group[0].replace('D', '')
        
        df_sim = pd.read_csv(similarity.farmclass.path)
                
        df_sim = df_sim[['experiment','compound','concentration',
                         'cellline','pval','padj',
                         'effect', 'brd', 'drug',
                         'time', 'conc', 'cl', 'gene'
                         ]]
        
        df_sim = df_sim[(df_sim['compound']==int(lsim_group[0])) &
                        (df_sim['concentration']==int(lsim_group[1])) & 
                        (df_sim['cellline']==lsim_group[2])]        
        
        df_sim.fillna(' ', inplace=True)
        
        df_sim.sort('padj', inplace=True)
        
        df_sim = df_sim.head(100)        
        
        #raise Exception('sim-sim')   
        
        output_json = df_sim.to_json(orient='values')
        response_data = {'data': json.loads(output_json)}
        
        return HttpResponse(json.dumps(response_data), content_type="application/json")

class ReportCorrelationTableJson(TemplateView): 
    template_name="report/report_detail.html"
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(ReportCorrelationTableJson, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        
        file_name = request.GET.get('file_name')            
            
        df_corr = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name)
                
        df_corr = df_corr[['sample','plate','correlation','series','platform','title']]        
        
        df_corr.fillna('No data', inplace=True)
        
        df_corr['plate'] = df_corr['plate'].str.replace('Tumour_', '')
        
        #raise Exception('corr') 
        
        output_json = df_corr.to_json(orient='values')
        response_data = {'data': json.loads(output_json)}
        
        return HttpResponse(json.dumps(response_data), content_type="application/json")    
    

class ReportPotentialTargetsJson(TemplateView): 
    template_name="report/report_detail.html"
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(ReportPotentialTargetsJson, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        report = Report.objects.get(pk=self.request.GET['reportID'])        
        
        pottargs = report.deeplearning_set.all()[2]
        pottargs_group_name = self.request.GET['pt_group_name']            
        
        
        lpottargs_group = pottargs_group_name.split('_')
        lpottargs_group[0] = lpottargs_group[0].replace('D', '')
        
        df_pottargs = pd.read_csv(pottargs.farmclass.path, sep='\t')
        
        column_name = 'SPS_'+lpottargs_group[0]+'_C_'+lpottargs_group[1]+'_celltype_'+lpottargs_group[2]+'.tab'# like SPS_10_C_100_celltype_A549.tab
                
        df_pottargs = df_pottargs[['Entrez ID','Gene Symbol', column_name]]
        
        #interset with genes df
        gene_group = GeneGroup.objects.get(report=report, name=self.request.GET['pt_group_name'])
        
        df_gene_group = pd.read_csv(gene_group.document.path, sep='\t', index_col='SYMBOL')
        
        df_pottargs.set_index('Gene Symbol', inplace=True)
        
        joinded_df = df_pottargs.join(df_gene_group, how='inner')
        
        joinded_df.index.name = 'Gene Symbol'
        joinded_df.reset_index(inplace=True)
        
        joinded_df = joinded_df[['Entrez ID','Gene Symbol', column_name]]        
        
        joinded_df.sort(column_name, inplace=True)
        
        joinded_df = joinded_df.head(100) 
               
        #raise Exception('stop')           
        
        #raise Exception('sim-sim')   
        
        output_json = joinded_df.to_json(orient='values')
        response_data = {'data': json.loads(output_json)}
        
        return HttpResponse(json.dumps(response_data), content_type="application/json")    
    
    
class ReportDSTableJson(TemplateView): 
    template_name="report/report_detail.html"
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(ReportDSTableJson, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        
        report = Report.objects.get(pk=request.GET.get('reportID'))
        file_name = request.GET.get('file_name')
        group_name = request.GET.get('ds_group_name')        
        
        try:    
            df_ds = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name+'_tblcut.csv',
                                  sep=None)
        except:
            df_ds = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name+'_tblcut.csv',
                                  sep='\t')
        #raise Exception('Drug Score')        
        df_ds.drop('Condition', axis=1, inplace=True)
        df_ds.fillna(1, inplace=True)
        #df_ds.Concentration.replace(to_replace=dict(''='NA'), inplace=True)                 
        df_ds.Concentration[df_ds.Concentration == '-666na'] = 'NA' 
         
        df_ds = df_ds[(df_ds['effect']>3)]         
        
        
        df_ds['padj'] = df_ds['padj'].map('{:,.2e}'.format)
        df_ds['pval'] = df_ds['pval'].map('{:,.2e}'.format)
        
        df_ds['padj'] = df_ds['padj'].apply(str)
        df_ds['pval'] = df_ds['pval'].apply(str)
        
        output_json = df_ds.to_json(orient='values')
        #raise Exception('ds') 
        response_data = {'data': json.loads(output_json)}
        
        return HttpResponse(json.dumps(response_data), content_type="application/json")    
    
class ReportDSBoxplotJson(TemplateView):
    
    template_name="report/report_detail.html"
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(ReportDSBoxplotJson, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        
        file_name  = request.GET.get('file_name')        
        report = Report.objects.get(pk=request.GET.get('reportID'))
        ds_group_name = request.GET.get('ds_group_name')
        pert_type = request.GET.get('pert_type')
        
        if pert_type == 'gene':
            try:    
                df_ds = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name+"_tblcut.csv",
                                  sep=None, index_col='pert_id')
            except:
                df_ds = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name+"_tblcut.csv",
                                  sep='\t', index_col='pert_id')
        
             
            
            #df_ds = df_ds[(df_ds['effect']>3)]
            #df_ds.reset_index(inplace=True)

            df_top10 = df_ds[['effect']]      
        
            df_top10 = df_top10.groupby(df_ds.index, level=0).mean() # take mean by effect column
        
            df_top10.sort(columns=['effect'], ascending=False, inplace=True)
        
        
            df_top10 = df_top10.filter(like='TRCN', axis=0)
                
        
            df_top10 = df_top10[:10] #take best 10 hits
        
            df_ds.reset_index(inplace=True)
        
            lbar_name = []
            lbar_val = [] 
        
            df_list = []
        
            for index, val in df_top10.iterrows(): 
                df_pert = df_ds[df_ds['pert_id'] == index]
                df_pert = df_pert[['Perturbation', 'pert_id', 'effect']]
                df_list.append(df_pert)
                df_pert.set_index('Perturbation', inplace=True)
            
                df_pert = df_pert.groupby(df_ds.index, level=0).mean()
            
                for index, val in df_pert.iterrows():
                    lbar_name.append(index)
                    lbar_val.append(val['effect'])
              
            response_data = {
                         'name': 'barplot',
                         'categories_name': lbar_name,
                         'series': [ round(elem, 2) for elem in lbar_val ],
                         
                         }        
        
        
            #raise Exception(' dsboxplot')
        
        
            return HttpResponse(json.dumps(response_data), content_type="application/json")
        
        
        if pert_type == 'molecule':
            try:    
                df_ds = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name+"_tblcut.csv",
                                  sep=None, index_col='pert_id')
            except:
                df_ds = pd.read_csv(settings.MEDIA_ROOT+"/"+file_name+"_tblcut.csv",
                                  sep='\t', index_col='pert_id')        
        
            df_ds = df_ds[(df_ds['effect']>3)]
            df_ds.reset_index(inplace=True)
            df_count = df_ds['pert_id'].value_counts()
        
            df_ds.set_index('pert_id', inplace=True )
            
            df_ds_contrib = df_ds[df_count>=3]
            df_ds_addit = df_ds[df_count<3]            
        
            #raise Exception('count')
            df_top10 = df_ds_contrib[['effect']]      
        
            df_top10 = df_top10.groupby(df_ds_contrib.index, level=0).mean() # take mean by effect column
        
            df_top10.sort(columns=['effect'], ascending=False, inplace=True)
        
            df_top10 = df_top10.filter(like='BRD-', axis=0)
            
            df_ds_addit = df_ds_addit[['effect']]      
            df_ds_addit = df_ds_addit.groupby(df_ds_addit.index, level=0).mean() # take mean by effect column
            df_ds_addit.sort(columns=['effect'], ascending=False, inplace=True)            
            df_ds_addit = df_ds_addit.filter(like='BRD-', axis=0)
        
        
            df_top10 = df_top10[:10] #take best 10 hits
            
            #if there is less than 10 hits fill the top10 with data from df_ds_addit
            n = len(df_top10.index)
            num_to_add = 10-n
            
            df_addit_to_add = df_ds_addit[:num_to_add]
            
            df_top10 = df_top10.append(df_addit_to_add)
            
            df_top10.sort(columns=['effect'], ascending=False, inplace=True)
        
            df_ds.reset_index(inplace=True)
        
            lbar_name = []
            lbar_val = [] 
        
            df_list = []
        
            for index, val in df_top10.iterrows(): 
                df_pert = df_ds[df_ds['pert_id'] == index]
                df_pert = df_pert[['Perturbation', 'pert_id', 'effect']]
                df_list.append(df_pert)            
        
            df_overal = pd.DataFrame()
            for df in df_list:
                df_overal = df_overal.append(df)
            
            #df_overal.reset_index(inplace=True)
            #df_overal.to_csv(settings.MEDIA_ROOT+"/"+file_name+"_cut_molecule1.csv")   
            
            series_data = []
            series_category = []
            
            for df in df_list:          
                
                series_category.append(df['Perturbation'].iloc[0])
                r_tumour = df.filter(like='effect')
                
                #
                
                median = np.log2(np.median(r_tumour)) 
                upper_quartile = np.log2(np.percentile(r_tumour, 75))
                lower_quartile = np.log2(np.percentile(r_tumour, 25))
                iqr = upper_quartile - lower_quartile
                upper_whisker = np.log2(r_tumour[r_tumour<=upper_quartile+1.5*iqr].max()[0])
                lower_whisker = np.log2(r_tumour[r_tumour>=lower_quartile-1.5*iqr].min()[0])
        
                if np.isnan(lower_whisker) or np.isinf(lower_whisker):
                    lower_whisker = lower_quartile
                if np.isnan(upper_whisker) or np.isinf(upper_whisker):
                    upper_whisker = upper_quartile
              
                lSerie = [lower_whisker, lower_quartile, median, upper_quartile, upper_whisker]
                #raise Exception('fuck')
                #lSerie[lSerie == -np.inf] = 0
                
                
                #lSerie.replace([np.inf, -np.inf], 0)
                 
                series_data.append(lSerie)
            
            response_data = {
                            'name': 'boxplot',
                            'categories': series_category,              
                            'data': series_data,
                            'tooltip': {
                            'headerFormat': '<em>Group: {point.key}</em><br/>'
                                       }
                             }
                     
        
        
            #raise Exception(' dsboxplot')
        
        
            return HttpResponse(json.dumps(response_data), content_type="application/json")
        
class ReportTargetInferenceJson(TemplateView):
    template_name="report/report_detail.html"
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(ReportTargetInferenceJson, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        
        report = Report.objects.get(pk=request.GET.get('reportID'))
        
        df_gene = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/TargetInferenceTable.csv')
        
        output_json = df_gene.to_json(orient='values')
        response_data = {'data': json.loads(output_json)}
        
        return HttpResponse(json.dumps(response_data), content_type="application/json")
                

class ReportMesenteryJson(TemplateView):
    template_name="report/report_detail.html"
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(ReportMesenteryJson, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        
        report = Report.objects.get(pk=request.GET.get('reportID'))
        file_name = request.GET.get('file_name')
        
        df_gene = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/'+file_name, sep=None)
        
        output_json = df_gene.to_json(orient='values')
        response_data = {'data': json.loads(output_json)}
        
        return HttpResponse(json.dumps(response_data), content_type="application/json") 


class ReportDeepGSKJson(TemplateView):
    template_name="report/report_detail.html"
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(ReportDeepGSKJson, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        
        report = Report.objects.get(pk=request.GET.get('reportID'))
        file_name = request.GET.get('file_name')
        
        df_gene = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/'+file_name)
        
        output_json = df_gene.to_json(orient='values')
        response_data = {'data': json.loads(output_json)}
        
        return HttpResponse(json.dumps(response_data), content_type="application/json")    
       
class ReportTest(TemplateView):
    
    """ Just testing playground to store useful scripts"""
    
    template_name="report/report_detail.html"
    
    def dispatch(self, request, *args, **kwargs):
        
        return super(ReportTest, self).dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        
        df = pd.read_csv("/home/mikhail/Downloads/GSK/GSK_divi_comparisons_update.csv" )
        
        df = df[df['PROTOCOL_DR_GROUP_TYPE']=='Medium']
        df = df[df['DURATION']==4] 
        df = df[df['MESENTERY_3']>0] 
        
        #df = df[df['TISSUE']!='Smooth Muscle (LCM Mesentery) RNA']
        
        df.to_csv("/home/mikhail/Downloads/GSK/GSK_4d_high.csv")
        raise Exception('fuck')
        
        import docx
        
        #########3 TARGET INFERENCE
        
        
        
        df = pd.read_csv("/home/mikhail/Downloads/GSK/GSK_divi_comparisons_update.csv", 
                           )
        cols = df.columns
        
        for col in cols:
            df1 = df[col]
            df_up = df1.copy()
            df_up.sort(ascending=False, inplace=True)
            df_up =  pd.DataFrame(df_up[:20])
            df_up.reset_index(inplace=True)
            
            paragraph = doc.add_heading(col, level=3)
            paragraph1 = doc.add_heading('Top 20 targets', level=4)
            
            t = doc.add_table(df_up.shape[0]+1, df_up.shape[1])            
            t.style = 'TableGrid'
            
            # add the header rows.
            for j in range(df_up.shape[-1]):
                t.cell(0,j).text = df_up.columns[j]

            # add the rest of the data frame
            for i in range(df_up.shape[0]):
                for j in range(df_up.shape[-1]):
                    t.cell(i+1,j).text = str(df_up.values[i,j])
            
            #raise Exception('cycle')
            
        doc.save('/home/mikhail/Downloads/GSK/WORD/Respiratory/tarinfer/tarinfer.docx')
        raise Exception('tarinfer')
        ##########3 PATHWAYS
        doc = docx.Document('/home/mikhail/Downloads/GSK/WORD/Respiratory/path/test.docx')
        
        df = pd.read_csv("/home/mikhail/Downloads/GSK/WORD/Respiratory/path/tbl-path_all.csv",
                            index_col=0, encoding='utf-8')
        cols = df.columns
        for col in cols:
            df1 = df[col]
            df_up = df1.copy()
            df_down = df1.copy()
            df_down.sort(ascending=True, inplace=True)
            df_up.sort(ascending=False, inplace=True)
            df_up =  pd.DataFrame(df_up[:20])
            df_down = pd.DataFrame(df_down[:20])
            
            df_up.reset_index(inplace=True)
            df_up.columns = ['Pathway', 'PAS']
            
            
            paragraph = doc.add_heading(col, level=3)
            paragraph1 = doc.add_heading('Top 20 up-regulated pathways', level=4)
            
            t = doc.add_table(df_up.shape[0]+1, df_up.shape[1])            
            t.style = 'TableGrid'
            # add the header rows.
            for j in range(df_up.shape[-1]):
                t.cell(0,j).text = df_up.columns[j]

            # add the rest of the data frame
            for i in range(df_up.shape[0]):
                for j in range(df_up.shape[-1]):
                    t.cell(i+1,j).text = str(df_up.values[i,j])
            
            ##############################
            df_down.reset_index(inplace=True)
            df_down.columns = ['Pathway', 'PAS']
            
            
            #paragraph = doc.add_paragraph(col)
            paragraph1 = doc.add_heading('Top 20 down-regulated pathways', level=4)
            
            t = doc.add_table(df_down.shape[0]+1, df_down.shape[1])            
            t.style = 'TableGrid'
            # add the header rows.
            for j in range(df_down.shape[-1]):
                t.cell(0,j).text = df_down.columns[j]

            # add the rest of the data frame
            for i in range(df_down.shape[0]):
                for j in range(df_down.shape[-1]):
                    t.cell(i+1,j).text = str(df_down.values[i,j])
            
            
            #raise Exception('stop')
        doc.save('/home/mikhail/Downloads/GSK/WORD/Respiratory/path/test.docx')
        raise Exception('Word')
        
        df = pd.read_csv("/home/mikhail/Downloads/AboveAndBeyond/Disease3S7/export_DE/Disease3S7.tab",
                            sep='\t', index_col=0)
        df.to_csv("/home/mikhail/Downloads/AboveAndBeyond/Disease3S7/export_DE/Disease3S7.csv",
                            sep=',')
        raise Exception('ramdom')
        
        df = pd.read_csv("/home/mikhail/Downloads/GSK/Database/AFFY_PROBE_SET.txt",
                           sep='\t', index_col=0)
        
        df_c = pd.read_csv("/home/mikhail/Downloads/GSK/4D/High/DE/Comparison74.tab",
                           sep='\t', index_col=0)
        
        df_c.to_csv('/home/mikhail/Downloads/GSK/4D/High/DE/Comparison74.tab', sep=',')
        raise Exception('fuck')
        import os
        
        ff = []
        for subdir, dirs, files in os.walk('/home/mikhail/Downloads/GSK/4D/High/PAS/'):
            for f in files:
                ff.append(f)
        
        ff.sort()
        
        for ffile in ff:
            df = pd.read_csv('/home/mikhail/Downloads/GSK/4D/High/PAS/'+ffile, sep=None, index_col=0)
            df.index.name = 'Pathway'
            df.to_csv('/home/mikhail/Downloads/GSK/4D/High/PAS/csv/'+ffile, sep=',')
            #raise Exception('cycle')
        
        raise Exception("GSK")
        
        
        lfiles = ['Comparison97',
        'Comparison91',
        'Comparison75',
        'Comparison74',
        'Comparison66',
        'Comparison6',
        'Comparison51',
        'Comparison42',
        'Comparison240',
        'Comparison234',
        'Comparison233',
        'Comparison228',
        'Comparison225',
        'Comparison224',
        'Comparison219',
        'Comparison21',
        'Comparison206',
        'Comparison203',
        'Comparison20',
        'Comparison191',
        'Comparison188',
        'Comparison179',
        'Comparison173',
        'Comparison160',
        'Comparison159',
        'Comparison151',
        'Comparison136',
        'Comparison127',
        'Comparison12',
        'Comparison106',
        'Comparison105'
]
        
        
        df_c = pd.read_csv("/home/mikhail/Downloads/GSK/4D/gsk_divi_comparisons_nocecdopa_sterilesodium_final.txt",
                           sep='\t', index_col=0)
        
        lout = []
        for comp in lfiles:
                       
            
            df_f = df_c[df_c['PROTOCOL_DR_GROUP_TYPE'].str.contains('High')]
            
            #df_f = df_c[df_c['comparison'].str.contains(comp)]
            
            df_f = df_f[df_f['MESENTERY_3']>0]
            
            df_f = df_f[df_f['DURATION']>1]
            
            df_f = df_f[df_f['TISSUE_SHORT']!='CEC']
            
            for index, row in df_f.iterrows():
                
                dout = {
                        'comparison': row['comparison'],
                        'TISSUE_SHORT': row['TISSUE_SHORT'],
                        'PROTOCOL_DR_GROUP_TYPE': row['PROTOCOL_DR_GROUP_TYPE'],
                        'EXPERIMENT_ID': row['EXPERIMENT_ID'],
                        'DURATION': row['DURATION'],
                        'COMPOUND_NAME': row['COMPOUND_NAME'] 
                        }
                lout.append(dout)
                
            df_f.sort(['COMPOUND_NAME', 'TISSUE_SHORT'], inplace=True)
            df_f.to_csv('/home/mikhail/Downloads/GSK/4D/prosecced_high_1.csv')
            #raise Exception('cycle')
        



        
        out = pd.DataFrame(lout)
        out.sort(['COMPOUND_NAME', 'TISSUE_SHORT'], inplace=True)
        #out.to_csv('/home/mikhail/Downloads/GSK/4D/prosecced_low.csv')
        raise Exception('stop')
        """
        lpert = ['Doxorubicin-high',
                 'GSK882380-high',
                 'GSK894281A-high',
                 'GSK1776357A-high',
                 'GW679769-high',
                 'GW693085-high']
        
        combinations= list(itertools.combinations(lpert, 4))
        
        
        ll = []
        for gg in combinations:
            sss = 'vs'.join(map(str,gg))
            ll.append(sss)
        
        group = ', '.join(ll)
        
        hhhh = len(group)
        raise Exception('perturb')
        """
        
        
        df = pd.read_csv("/home/mikhail/Downloads/AboveAndBeyond/s6s7/DE_Disease4_S6.tab", sep='\t')
        
        df.to_csv("/home/mikhail/Downloads/AboveAndBeyond/s6s7/DE_Disease4_S6.csv", sep=',', index=False)
        raise Exception('stop')
        """
        # Cut table Effect >3
        df = df[df.effect>3]
        
        df.to_csv("/home/mikhail/Downloads/AboveAndBeyond/res_DS_Disease4s7_final.txt_tblcut.csv", index=False)
        """
        
        df = df[['Perturbation', 'pert_id', 'effect']]
        
        
        raise Exception('stop')
        response_data  = {
                          'result': True
                          }
        return HttpResponse(json.dumps(response_data), content_type="application/json")    
    
    
    