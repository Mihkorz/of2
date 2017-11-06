import json
import pandas as pd
import csv
import numpy as np
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches

from django.views.generic.base import TemplateView
from django.views.generic.detail import DetailView
from django.conf import settings

from .models import Report, GeneGroup, PathwayGroup, TfGroup
from core.models import Pathway, Node, Component


class ReportGeneratePDF(DetailView):
    """
    Generate PDF Reports
    Originaly made gro gsk_prj4 reports
    """
    
    model = Report
    template_name = 'report/generate_pdf.html'
    
    tox_paths = ['reactome_Toxicity_of_botulinum_toxin_type_B_BoNT_B__Main_Pathway',
                        'reactome_Toxicity_of_botulinum_toxin_type_G_BoNT_G__Main_Pathway',
                        'reactome_Toxicity_of_botulinum_toxin_type_C_BoNT_C__Main_Pathway',
                        'reactome_Toxicity_of_botulinum_toxin_type_E_BoNT_E__Main_Pathway',
                        'reactome_Toxicity_of_botulinum_toxin_type_A_BoNT_A__Main_Pathway',
                        'reactome_Toxicity_of_botulinum_toxin_type_D_BoNT_D__Main_Pathway',
                        'reactome_Toxicity_of_botulinum_toxin_type_F_BoNT_F__Main_Pathway',
                        'reactome_Negative_regulators_of_RIG_I_MDA5_signaling_Main_Pathway',
                        'NCI_Cellular_roles_of_Anthrax_toxin_Pathway_.MAPKKK_cascade.',
                        'NCI_Cellular_roles_of_Anthrax_toxin_Pathway_.monocyte_activation.',
                        'NCI_Cellular_roles_of_Anthrax_toxin_Pathway_.apoptosis.',
                        'NCI_Cellular_roles_of_Anthrax_toxin_Pathway_.inflammatory_response.',
                        'NCI_Cellular_roles_of_Anthrax_toxin_Pathway_.necrosis.',
                        'NCI_Cellular_roles_of_Anthrax_toxin_Pathway_.negative_regulation_of_macrophage_activation.',
                        'NCI_Cellular_roles_of_Anthrax_toxin_Pathway_.negative_regulation_of_myeloid_dendritic_cell_antigen_processing_and_presentation.',
                        'NCI_Cellular_roles_of_Anthrax_toxin_Pathway_.negative_regulation_of_phagocytosis.',
                        'NCI_Cellular_roles_of_Anthrax_toxin_Pathway_.platelet_activation.',
                        'NCI_Cellular_roles_of_Anthrax_toxin_Pathway_.regulation_of_endothelial_cell_proliferation.',
                        'NCI_Cellular_roles_of_Anthrax_toxin_Pathway_.macrophage_activation.',
                        'NCI_IL2_mediated_signaling_events_Pathway_.natural_killer_cell_mediated_cytotoxicity.',
                        'NCI_IL23_mediated_signaling_events_Pathway_.positive_regulation_of_T_cell_mediated_cytotoxicity.',
                        'NCI_IL27_mediated_signaling_events_Pathway_.positive_regulation_of_T_cell_mediated_cytotoxicity.',
                        'NCI_IL12_signaling_mediated_by_STAT4_Pathway_.natural_killer_cell_mediated_cytotoxicity.',
                        'biocarta_ras_independent_pathway_in_nk_cell_mediated_cytotoxicity_Pathway_.positive_regulation_of_natural_killer_cell_mediated_cytotoxicity.',
                        'biocarta_ras_independent_pathway_in_nk_cell_mediated_cytotoxicity_Main_Pathway',
                        'KEGG_Natural_killer_cell_mediated_cytotoxicity_Main_Pathway',
                        'KEGG_Graft_versus_host_disease_Main_Pathway',
                        'KEGG_Huntington_s_disease_Main_Pathway',
                        'KEGG_Vibrio_cholerae_infection_Main_Pathway',
                        'KEGG_Chemical_carcinogenesis_Main_Pathway',
                        'KEGG_Allograft_rejection_Main_Pathway',
                        'biocarta_apoptotic_signaling_in_response_to_dna_damage_Main_Pathway',
                        'biocarta_apoptotic_signaling_in_response_to_dna_damage_Pathway_.apoptosis.',
                        'biocarta_mechanism_of_acetaminophen_activity_and_toxicity_Main_Pathway',
                        'biocarta_mechanism_of_acetaminophen_activity_and_toxicity_Pathway_.fever.',
                        'biocarta_mechanism_of_acetaminophen_activity_and_toxicity_Pathway_.inflammatory_response.',
                        'biocarta_mechanism_of_acetaminophen_activity_and_toxicity_Pathway_.platelet_activation.',
                        'biocarta_mechanism_of_acetaminophen_activity_and_toxicity_Pathway_.prostaglandin_biosynthetic_process.',
                        'reactome_Uptake_and_function_of_anthrax_toxins_Main_Pathway',
                        'reactome_Oxidative_Stress_Induced_Senescence_Main_Pathway',
                        'mTOR_Pathway_Metabolism_Stress_Response_and_Apoptosis',
                        'reactome_DNA_Damage_Telomere_Stress_Induced_Senescence_Main_Pathway',
                        'Cellular_Apoptosis_Pathway',
                        'Mitochondrial_Apoptosis_Pathway',
                        'ATM_Pathway_Apoptosis',
                        'ATM_Pathway_Apoptosis_and_Senescence',
                        'Mitochondrial_Apoptosis_Pathway_Depolarization',
                        'NCI_ATM_Pathway_.apoptosis.',
                        'ceramide_ide_novoi_biosynthesis',
                        'NCI_TNF_receptor_signaling_pathway__Main_Pathway',
                        'TNF_Signaling_Pathway',
                        'TNF_Signaling_Pathway_Apoptosis',
                        'KEGG_TNF_signaling_Main_Pathway',
                        'reactome_CYP2E1_reactions_Main_Pathway'
                        ]
    
    
    def dispatch(self, request, *args, **kwargs):
        return super(ReportGeneratePDF, self).dispatch(request, *args, **kwargs)
    
    def get_context_data(self, **kwargs):
              
        context = super(ReportGeneratePDF, self).get_context_data(**kwargs)
        
        report = self.object
        
        document = Document()
        
        """
        PATHWAY LEVEL
        """ 
        
        head = document.add_heading("Pathway level analysis", 2)
        head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph()
        p = document.add_paragraph('Here we collected all BioMAP samples in complete diversity set and all mapped biomarkers to corresponding genes, then for each biomarker we calculated the ratio "Log Ratio"/"Significance prediction envelope at 95%" aggregated", after that for genes present across several assays we took the mean of that ratio and mapped the resulting values onto the pathway database.')
        p = document.add_paragraph()
        
        
            
        p = document.add_paragraph()
        p = document.add_heading('BioMAP diversity data', level=3)
        p = document.add_paragraph() 
            
        file_name = 'biomap_diversity.pathway_scores.csv' 
        df_path = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/'+file_name, sep=None)
        try:
            df_path.columns = ['Pathway', 'PAS']
        except:
            df_path.columns = ['Pathway', 'N6948-21-1_20000nM', 'N12721-44-A1_20000nM']
        
        df_output = df_path        
        
            
            
        df_up = df_output.copy()
        df_down = df_output.copy()
        df_tox = df_output.copy()
        df_down.sort_values(by="PAS", ascending=True, inplace=True)
        df_up.sort_values(by="PAS", ascending=False, inplace=True)
        df_up =  pd.DataFrame(df_up[:20])
        df_down = pd.DataFrame(df_down[:20])
            
        df_tox = df_tox.ix[self.tox_paths] 
            
            
        
            
        paragraph1 = document.add_heading('Top 20 up-regulated pathways', level=4)
            
        t = document.add_table(df_up.shape[0]+1, df_up.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df_up.shape[-1]):
            t.cell(0,j).text = df_up.columns[j]

        # add the rest of the data frame
        for i in range(df_up.shape[0]):
            for j in range(df_up.shape[-1]):
                t.cell(i+1,j).text = str(df_up.values[i,j])
            
        ############################## DOWN
        
            
            
        #paragraph = doc.add_paragraph(col)
        paragraph1 = document.add_heading('Top 20 down-regulated pathways', level=4)
            
        t = document.add_table(df_down.shape[0]+1, df_down.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df_down.shape[-1]):
            t.cell(0,j).text = df_down.columns[j]

        # add the rest of the data frame
        for i in range(df_down.shape[0]):
            for j in range(df_down.shape[-1]):
                t.cell(i+1,j).text = str(df_down.values[i,j])
                    
         
        head = document.add_heading("Toxcast analysis", 2)
        head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph()
        p = document.add_paragraph('We collected ToxCast data for SAg and 3C systems from EPA. For SAg we found 10 common biomarkers that are present in investigated dataset, these include: CD38, CD40, CD69, CD62E/E-Selectin, CXCL8/IL-8, CCL2/MCP-1, CXCL9/MIG, PBMC Cytotoxicity, Proliferation, SRB. For 3C we found 12 common biomarkers: CD62E/E-Selectin, CXCL8/IL-8, CCL2/MCP-1, CXCL9/MIG, PBMC Cytotoxicity, Proliferation, SRB, CD54/ICAM-1, HLA-DR, CD141/Thrombomodulin, CD142/Tissue Factor, CD106/VCAM-1, CD87/uPAR. For each drug perturbation from BioMAP we correlated log-ratio/conf.envelope to downward curve-fit analysis readout from ToxCast. Results can be found in two following tables.')
        p = document.add_paragraph('Input:\n Biomap datatable (Download file):\n biomarkers as rows, perturbations as columns; ToxCast tables with readouts for 3C \n and SAg: \n biomarkers as columns, perturbations as rows.')
        p = document.add_paragraph('Processing: Within each of the two systems (3C, SAg) for each compound perturbation pair (GSKvsToxCast) rows (biomarkers) from BioMAP datatable were matched to columns (biomarkers) of ToxCast table. For each compound pair (GSKvsToxCast) this resulted in two vectors of length 10 (in case of SAg system) and 12 (in case of 3C system). Pearson correlation was calculated for all possible compound perturbation pairs (GSKvsToxCast).')
        p = document.add_paragraph('Output: A table (with ToxCast perturbations as rows and BioMAP perturbations as columns) populated with Pearson correlation values. As the calculations for 3C and SAg were done independently, two separate tables were generated.')
        
        
        file_name = 'granular_toxcast_3C.txt' 
        df_3c = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/'+file_name, sep='\t')
        file_name = 'granular_toxcast_SAg.txt' 
        df_sag = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/'+file_name, sep='\t')
        
        df_3c.fillna(0, inplace=True)
        df_3c = df_3c.head(100)
        df_val = df_3c.iloc[:,2:]
        df_val= df_val.round(decimals=2)
        
        df_output = df_3c.iloc[:, :2]
        df_output = pd.concat([df_output, df_val], axis=1)
        
        p = document.add_paragraph()
        p = document.add_heading('3C results', level=4)
        p = document.add_paragraph()
        
        t = document.add_table(df_output.shape[0]+1, df_output.shape[1])   
                 
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df_output.shape[-1]):
            print "table head "+ str(j)
            t.cell(0,j).text = df_output.columns[j]

        # add the rest of the data frame
        for i in range(df_output.shape[0]):
            print "table body "+ str(i)
            for j in range(df_output.shape[-1]):
                t.cell(i+1,j).text = str(df_output.values[i,j])
        
        
        
        df_sag.fillna(0, inplace=True)
        df_sag = df_sag.head(100)
        df_val = df_sag.iloc[:,2:]
        df_val= df_val.round(decimals=2)
        
        df_output = df_sag.iloc[:, :2]
        df_output = pd.concat([df_output, df_val], axis=1)
        
        p = document.add_paragraph()
        p = document.add_heading('SAg results', level=4)
        p = document.add_paragraph()
        
        t = document.add_table(df_output.shape[0]+1, df_output.shape[1])   
                 
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df_output.shape[-1]):
            print "table head "+ str(j)
            t.cell(0,j).text = df_output.columns[j]

        # add the rest of the data frame
        for i in range(df_output.shape[0]):
            print "table body "+ str(i)
            for j in range(df_output.shape[-1]):
                t.cell(i+1,j).text = str(df_output.values[i,j])  
        
        document.save(settings.MEDIA_ROOT+'/report-pdf/'+report.slug+'.docx')
        print "path biomap done"
        
        raise Exception('GENE PATH')
        
        """
         Similarity to LINCS dataset
        """
        
        head = document.add_heading("Similarity to LINCS dataset", 2)
        head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph()
        p = document.add_paragraph('In this section we calculated the similarity of investigated perturbations to LINCS dataset. For each sample we performed the correlation analysis of its gene expression values and 1.3+ million samples from LINCS on the probe level. Then for each sample we picked top-200 LINCS samples by correlation coefficient and combined these lists in the resulting table. Within each perturbation we averaged correlation coefficients between two replicates. In LINCS dataset pert_ids with "BRD-", "TRCN-" and "BRDN-" prefixes designate small molecule, knock-down and ORF overexpression perturbations, respectively. In pert_desc column you can find perturbation name (name of the molecule or name of the gene in case of knock-downs and ORF overexpression).')
        p = document.add_paragraph()
        
        
        
        file_name = 'q2norm/correlation_top3.csv'
        df_sim = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/'+file_name)
        df_sim = df_sim.head(80)
        df_sim.fillna(0, inplace=True)
        df_val = df_sim.iloc[:,6:]
        df_val= df_val.round(decimals=2)
        
        n=10 # break DF in chunks of 10 columns
        
        k = len(list(df_val.columns))%n
        list_df = [df_val.iloc[:,i*n:i*n+n] for i in range(0,k)]
        
        
        
        for idx, df in enumerate(list_df):
            print idx
            p = document.add_paragraph()
            p = document.add_heading('Similarity scoring of GSK L1000 data in comparison to LINCS dataset. Part '+str(idx+1), level=4)
            p = document.add_paragraph()
            
            df_output = df_sim.iloc[:, :6]
            df_output = pd.concat([df_output, df], axis=1)
            if idx == 2:
                raise Exception('haha')
            
            t = document.add_table(df_output.shape[0]+1, df_output.shape[1])            
            t.style = 'TableGrid'
            # add the header rows.
            for j in range(df_output.shape[-1]):
                print "table head "+ str(j)
                t.cell(0,j).text = df_output.columns[j]

            # add the rest of the data frame
            for i in range(df_output.shape[0]):
                print "table body "+ str(i)
                for j in range(df_output.shape[-1]):
                        t.cell(i+1,j).text = str(df_output.values[i,j])
        
        
        document.save(settings.MEDIA_ROOT+'/report-pdf/Similarity/'+report.slug+'.docx')
        raise Exception('haha')
        
        
        """
        Hit qualification
        """
        head = document.add_heading("Hit qualification", 2)
        head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph()
        
        p = document.add_paragraph('Here we utilised our already established methodology to estimate the probability of failure in clinical trials (Artemov et al., 2016). On Figures 1 and 2 you can see the overall statistic for provided L1000 dataset. In "Hit qualification table" you can see the actual score of drug perturbations analysed in this report. ')
        
        document.add_picture(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t1/hit1.png', width=Inches(5.0))
        p = document.add_paragraph('Figure 1. Probability of success for drug perturbations from provided L1000 dataset.')
        document.add_picture(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t1/hit2.png', width=Inches(5.0))
        p = document.add_paragraph('Figure 2. Probability of failure for drug perturbations from provided L1000 dataset.')
        head = document.add_heading("Hit qualification", 4)
        
        file_name = 'clinical_trial/predCT.csv'
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/'+file_name)
        
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
                
        
        document.save(settings.MEDIA_ROOT+'/report-pdf/Hit/'+report.slug+'.docx')
        print "hit done"
        
        raise Exception('HIT')
        # End of HIT        
         
        """
        Header
        """
        head = document.add_heading(report.title, 0)
        head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        paragraph = document.add_paragraph('Organization: \nGSK')
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        paragraph = document.add_paragraph()
        paragraph = document.add_paragraph()
        paragraph = document.add_paragraph()
        paragraph = document.add_paragraph()
        paragraph = document.add_paragraph()
        paragraph = document.add_paragraph()
        
        paragraph = document.add_paragraph('Table of contents: \n \
                Gene level analysis \n \
                Pathway Analysis \n \
                Drug Repurposing \n \
                BioMAP inference \n \
                L1000-BioMAP correlation \n \
                Hit qualification \n \
                Tox groups analysis \n \
                Toxcast analysis \n \
                Similarity to LINCS dataset \n \
                Barcodes for L1000 data \n \
                Discussion \n ')
        
        p = document.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
    
        
        
        
        """
             GENE LEVEL
        """
        
        head = document.add_heading("Gene level analysis", 2)
        head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph()
        p = document.add_paragraph("Provided samples were distributed into comparison groups. Each comparison included stimulated or perturbed data and unstimulated controls. For gene level analysis each comparison was analyzed by limma algorithm.")
        p = document.add_paragraph()
        
        
        for gene_group in self.object.genegroup_set.all():
            p = document.add_paragraph()
            p = document.add_paragraph(gene_group.name+" vs Normal")
            p = document.add_paragraph() 
            
            df_gene = pd.read_csv(gene_group.doc_logfc)
            df_gene = df_gene[['SYMBOL', 'logFC', 'adj.P.Val']]
            
            df_gene = df_gene[(df_gene['adj.P.Val']<report.pval_theshold_plot) & (np.absolute(df_gene['logFC'])>report.logcf_theshold_plot)]
            
            df_gene['logFC'] = df_gene['logFC'].round(decimals=2)
            df_gene['adj.P.Val'] = df_gene['adj.P.Val'].map('{:,.2e}'.format)
            df_gene['adj.P.Val'] = df_gene['adj.P.Val'].apply(str)
            
            df_up = df_gene.copy()
            df_down = df_gene.copy()
            df_down.sort_values(by=['logFC'], ascending=True, inplace=True)
            df_up.sort_values(by=['logFC'], ascending=False, inplace=True)
            df_up =  pd.DataFrame(df_up[:20])
            df_down = pd.DataFrame(df_down[:20])
            
            df_output = df_up.append(df_down)
            df_output.drop_duplicates(inplace=True)
            
            
            document.add_picture(settings.MEDIA_ROOT+'/report-pdf/oncoPNG/'+report.slug+'/'+gene_group.name+'.png', width=Inches(5.0))
            
            p = document.add_heading('Top 20 up-regulated and down-regulated genes', level=4)
            
            t = document.add_table(df_output.shape[0]+1, df_output.shape[1])            
            t.style = 'TableGrid'
            # add the header rows.
            for j in range(df_output.shape[-1]):
                t.cell(0,j).text = df_output.columns[j]

            # add the rest of the data frame
            for i in range(df_output.shape[0]):
                for j in range(df_output.shape[-1]):
                    t.cell(i+1,j).text = str(df_output.values[i,j])
            
            #raise Exception(gene_group)
        
        # End of GENE LEVEL
        
        print "Gene done"
        
        """
        PATHWAY LEVEL
        """ 
        
        head = document.add_heading("Pathway level analysis", 2)
        head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph()
        p = document.add_paragraph("Pathway level analysis was performed on several curated pathway databases (KEGG, NCI, Reactome, SABiosciences) comprising 3117 signaling and metabolic pathways in total. To estimate pathway activation we used iPANDA algorithm.")
        p = document.add_paragraph()
        
        for path_group in self.object.pathwaygroup_set.all():
            
            p = document.add_paragraph()
            p = document.add_heading(path_group.name, level=3)
            p = document.add_paragraph() 
            
            df_path = pd.read_csv(path_group.document, sep='\t', index_col='Pathway')
            
            df1_tumour = df_path[[x for x in df_path.columns if 'Tumour' in x]]
            s1_tumour = df1_tumour.mean(axis=1).round(decimals=2)
            
            df_output = pd.DataFrame()        
            df_output['PAS'] = s1_tumour
            
            
            df_up = df_output.copy()
            df_down = df_output.copy()
            df_tox = df_output.copy()
            df_down.sort_values(by="PAS", ascending=True, inplace=True)
            df_up.sort_values(by="PAS", ascending=False, inplace=True)
            df_up =  pd.DataFrame(df_up[:20])
            df_down = pd.DataFrame(df_down[:20])
            
            df_tox = df_tox.ix[self.tox_paths] 
            
            
            df_up.reset_index(inplace=True)
            df_up.columns = ['Pathway', 'PAS']
            
            paragraph1 = document.add_heading('Top 20 up-regulated pathways', level=4)
            
            t = document.add_table(df_up.shape[0]+1, df_up.shape[1])            
            t.style = 'TableGrid'
            # add the header rows.
            for j in range(df_up.shape[-1]):
                t.cell(0,j).text = df_up.columns[j]

            # add the rest of the data frame
            for i in range(df_up.shape[0]):
                for j in range(df_up.shape[-1]):
                    t.cell(i+1,j).text = str(df_up.values[i,j])
            
            ############################## DOWN
            df_down.reset_index(inplace=True)
            df_down.columns = ['Pathway', 'PAS']
            
            
            #paragraph = doc.add_paragraph(col)
            paragraph1 = document.add_heading('Top 20 down-regulated pathways', level=4)
            
            t = document.add_table(df_down.shape[0]+1, df_down.shape[1])            
            t.style = 'TableGrid'
            # add the header rows.
            for j in range(df_down.shape[-1]):
                t.cell(0,j).text = df_down.columns[j]

            # add the rest of the data frame
            for i in range(df_down.shape[0]):
                for j in range(df_down.shape[-1]):
                    t.cell(i+1,j).text = str(df_down.values[i,j])
                    
            ############################## TOX
            df_tox.reset_index(inplace=True)
            df_tox.columns = ['Pathway', 'PAS']
            
            
            #paragraph = doc.add_paragraph(col)
            paragraph1 = document.add_heading('Tox signaling pathways', level=4)
            
            t = document.add_table(df_tox.shape[0]+1, df_tox.shape[1])            
            t.style = 'TableGrid'
            # add the header rows.
            for j in range(df_tox.shape[-1]):
                t.cell(0,j).text = df_tox.columns[j]

            # add the rest of the data frame
            for i in range(df_tox.shape[0]):
                for j in range(df_tox.shape[-1]):
                    t.cell(i+1,j).text = str(df_tox.values[i,j])
            
            #raise Exception('path stop')
        
        document.save(settings.MEDIA_ROOT+'/report-pdf/'+report.slug+'.docx')
        print "path biomap done"
        
        raise Exception('GENE PATH')
        print "path done"
        ################################ BioMAP diversity data
        
        head = document.add_heading("Pathway level analysis. BioMAP diversity data", 2)
        head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph()
        p = document.add_paragraph('Here we collected all BioMAP samples in complete diversity set and all mapped biomarkers to corresponding genes, then for each biomarker we calculated the ratio "Log Ratio"/"Significance prediction envelope at 95%" aggregated", after that for genes present across several assays we took the mean of that ratio and mapped the resulting values onto the pathway database.')
        p = document.add_paragraph()
        
        file_name = 'biomap_diversity.pathway_scores.csv'
        df_path = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/'+file_name, index_col='PATHWAY', encoding='utf-8')
        df_path.fillna(0, inplace=True)
        
        
        
        for path_group in df_path.columns:
            
            df_path1 = pd.DataFrame(df_path[path_group])
            df_path1.columns = [['PAS']]
            
            p = document.add_paragraph()
            p = document.add_heading(path_group, level=3)
            p = document.add_paragraph()
            
            df_up = df_path1.copy()
            df_down = df_path1.copy()
            df_down.sort_values(by="PAS", ascending=True, inplace=True)
            df_up.sort_values(by="PAS", ascending=False, inplace=True)
            df_up =  pd.DataFrame(df_up[:20])
            df_down = pd.DataFrame(df_down[:20])

            df_up.reset_index(inplace=True)
            df_up.columns = ['Pathway', 'PAS']
            df_up['Pathway'] = df_up['Pathway'].map(lambda x: x.encode('unicode-escape').decode('utf-8'))
            
            paragraph1 = document.add_heading('Top 20 up-regulated pathways', level=4)
            
            t = document.add_table(df_up.shape[0]+1, df_up.shape[1])            
            t.style = 'TableGrid'
            # add the header rows.
            for j in range(df_up.shape[-1]):
                t.cell(0,j).text = df_up.columns[j]

            # add the rest of the data frame
            for i in range(df_up.shape[0]):
                for j in range(df_up.shape[-1]):
                    t.cell(i+1,j).text = str(df_up.values[i,j])
            
            ############################## DOWN
            df_down.reset_index(inplace=True)
            df_down.columns = ['Pathway', 'PAS']
            df_down['Pathway'] = df_down['Pathway'].map(lambda x: x.encode('unicode-escape').decode('utf-8'))
            
            
            #paragraph = doc.add_paragraph(col)
            paragraph1 = document.add_heading('Top 20 down-regulated pathways', level=4)
            
            t = document.add_table(df_down.shape[0]+1, df_down.shape[1])            
            t.style = 'TableGrid'
            # add the header rows.
            for j in range(df_down.shape[-1]):
                t.cell(0,j).text = df_down.columns[j]

            # add the rest of the data frame
            for i in range(df_down.shape[0]):
                for j in range(df_down.shape[-1]):
                    t.cell(i+1,j).text = str(df_down.values[i,j]) 

        #end of BioMAP diversity data
        # END OF PATHWAY
        
        
        
        
        """
        DL
        """
        
        deeplearning = report.deeplearning_set.all()[0]
        df_farm = pd.read_csv(deeplearning.farmclass.path, index_col='Name', sep=None)
        """
        Drug Repurposing Therapeutic use
        """
        head = document.add_heading("Drug Repurposing", 2)
        head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph()
        
        for gene_group in self.object.genegroup_set.all():
            
            df_farm1 = df_farm[df_farm.fileName== gene_group.name]
            s_features = df_farm1.iloc[:, 9:]
            s_features = s_features.mean()         
            s_features.sort(ascending=False)
            
            s_features = s_features[s_features>0] 
            
            df_output = pd.DataFrame(s_features)
            df_output.reset_index(inplace=True)
            df_output.columns = ['Pharm class', 'Value']
            
            
            
            p = document.add_paragraph()
            p = document.add_heading('Therapeutic use for '+gene_group.name, level=4)
            p = document.add_paragraph()
        
            t = document.add_table(df_output.shape[0]+1, df_output.shape[1])   
                 
            t.style = 'TableGrid'
            # add the header rows.
            for j in range(df_output.shape[-1]):
                print "table head "+ str(j)
                t.cell(0,j).text = df_output.columns[j]

                # add the rest of the data frame
                for i in range(df_output.shape[0]):
                    print "table body "+ str(i)
                    for j in range(df_output.shape[-1]):
                        t.cell(i+1,j).text = str(df_output.values[i,j])
            
         
        
        p = document.add_paragraph()
        p = document.add_paragraph()
        
        head = document.add_heading("Top diseases", 4)
        
        file_name = 'granular_ds.txt'
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/'+file_name, sep='\t')
        
        df.sort_values(by="Score", ascending=False, inplace=True)
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])   
        
        document.save(settings.MEDIA_ROOT+'/report-pdf/Drug/'+report.slug+'.docx')
        raise Exception('toxgroup')
        
        """
        Tox groups analysis
        """
        head = document.add_heading("Similarity to LINCS dataset", 2)
        head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph()
        p = document.add_paragraph("Here we investigated the common and differential features among 5 Tox groups highlighted in the dataset annotation table. All samples corresponding to drugs from a particular group (case) were aggregated and compared to 'Rat clean' group (reference) that corresponds to perturbations with compounds that didn't demonstrate toxicity. Resulting lists of differentially expressed genes were compared across all 5 Tox groups and obrained overlaps are depicted below and listed in the table.")
        p = document.add_paragraph()
        
        document.add_picture(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t5/img_tox.png', width=Inches(5.0))
        p = document.add_paragraph('Differentially expressed genes distribution across 5 distinct Tox groups')
        
        file_name = 'gsk_prj4_t5/de_genes_binary.csv'
        df_output = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+file_name)
        
        p = document.add_paragraph()
        p = document.add_heading('Tox groups analysis', level=4)
        p = document.add_paragraph()
        
        t = document.add_table(df_output.shape[0]+1, df_output.shape[1])   
                 
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df_output.shape[-1]):
            print "table head "+ str(j)
            t.cell(0,j).text = df_output.columns[j]

        # add the rest of the data frame
        for i in range(df_output.shape[0]):
            print "table body "+ str(i)
            for j in range(df_output.shape[-1]):
                t.cell(i+1,j).text = str(df_output.values[i,j])
        
        document.save(settings.MEDIA_ROOT+'/report-pdf/Tox_groups_analysis.docx')
        raise Exception('toxgroup')
        
            
        #Similarity to LINCS dataset
        
        head = document.add_heading("Toxcast analysis", 2)
        head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph()
        p = document.add_paragraph('We collected ToxCast data for SAg and 3C systems from EPA. For SAg we found 10 common biomarkers that are present in investigated dataset, these include: CD38, CD40, CD69, CD62E/E-Selectin, CXCL8/IL-8, CCL2/MCP-1, CXCL9/MIG, PBMC Cytotoxicity, Proliferation, SRB. For 3C we found 12 common biomarkers: CD62E/E-Selectin, CXCL8/IL-8, CCL2/MCP-1, CXCL9/MIG, PBMC Cytotoxicity, Proliferation, SRB, CD54/ICAM-1, HLA-DR, CD141/Thrombomodulin, CD142/Tissue Factor, CD106/VCAM-1, CD87/uPAR. For each drug perturbation from BioMAP we correlated log-ratio/conf.envelope to downward curve-fit analysis readout from ToxCast. Results can be found in two following tables.')
        p = document.add_paragraph('Input:\n Biomap datatable (Download file):\n biomarkers as rows, perturbations as columns; ToxCast tables with readouts for 3C \n and SAg: \n biomarkers as columns, perturbations as rows.')
        p = document.add_paragraph('Processing: Within each of the two systems (3C, SAg) for each compound perturbation pair (GSKvsToxCast) rows (biomarkers) from BioMAP datatable were matched to columns (biomarkers) of ToxCast table. For each compound pair (GSKvsToxCast) this resulted in two vectors of length 10 (in case of SAg system) and 12 (in case of 3C system). Pearson correlation was calculated for all possible compound perturbation pairs (GSKvsToxCast).')
        p = document.add_paragraph('Output: A table (with ToxCast perturbations as rows and BioMAP perturbations as columns) populated with Pearson correlation values. As the calculations for 3C and SAg were done independently, two separate tables were generated.')
        
        
        file_name = 'granular_toxcast_3C.txt' 
        df_3c = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/'+file_name, sep='\t')
        file_name = 'granular_toxcast_SAg.txt' 
        df_sag = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/'+file_name, sep='\t')
        
        df_3c.fillna(0, inplace=True)
        df_3c = df_3c.head(100)
        df_val = df_3c.iloc[:,2:]
        df_val= df_val.round(decimals=2)
        
        df_output = df_3c.iloc[:, :2]
        df_output = pd.concat([df_output, df_val], axis=1)
        
        p = document.add_paragraph()
        p = document.add_heading('3C results', level=4)
        p = document.add_paragraph()
        
        t = document.add_table(df_output.shape[0]+1, df_output.shape[1])   
                 
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df_output.shape[-1]):
            print "table head "+ str(j)
            t.cell(0,j).text = df_output.columns[j]

        # add the rest of the data frame
        for i in range(df_output.shape[0]):
            print "table body "+ str(i)
            for j in range(df_output.shape[-1]):
                t.cell(i+1,j).text = str(df_output.values[i,j])
        
        
        
        df_sag.fillna(0, inplace=True)
        df_sag = df_sag.head(100)
        df_val = df_sag.iloc[:,2:]
        df_val= df_val.round(decimals=2)
        
        df_output = df_sag.iloc[:, :2]
        df_output = pd.concat([df_output, df_val], axis=1)
        
        p = document.add_paragraph()
        p = document.add_heading('SAg results', level=4)
        p = document.add_paragraph()
        
        t = document.add_table(df_output.shape[0]+1, df_output.shape[1])   
                 
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df_output.shape[-1]):
            print "table head "+ str(j)
            t.cell(0,j).text = df_output.columns[j]

        # add the rest of the data frame
        for i in range(df_output.shape[0]):
            print "table body "+ str(i)
            for j in range(df_output.shape[-1]):
                t.cell(i+1,j).text = str(df_output.values[i,j])
        
        
        
        document.save(settings.MEDIA_ROOT+'/report-pdf/Toxcastanalysis/'+report.slug+'.docx')
        raise Exception('stop')
       
        
        
        
        
        
       
        
        """
        Drug Repurposing
        """
        
        head = document.add_heading("Drug Repurposing", 2)
        head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph()
        p = document.add_paragraph()
        
        head = document.add_heading("Top diseases", 4)
        
        file_name = 'granular_ds.txt'
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/'+file_name, sep='\t')
        
        df.sort_values(by="Score", ascending=False, inplace=True)
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
                
        print "repurpose done"
        
        """
        BioMAP inference
        """ 
        """
        head = document.add_heading("BioMAP inference", 2)
        head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph()
        
        
        
        document.add_picture(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t5/g_distribution/gsk_aring.png', width=Inches(5.0))
        p = document.add_paragraph('Distribution of ARING property')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_picture(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t5/g_distribution/gsk_clogp.png', width=Inches(5.0))
        p = document.add_paragraph('Distribution of CLOGP property')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_picture(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t5/g_distribution/gsk_mw.png', width=Inches(5.0))
        p = document.add_paragraph('Distribution of MW property')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_picture(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t5/g_distribution/gsk_tpsa.png', width=Inches(5.0))
        p = document.add_paragraph('Distribution of TPSA property')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        
        p = document.add_paragraph('In this section we present the results of the experiment aimed at inferring BioMAP cytotox and diversity biomarkers independently from L1000 dataset using drug perturbation response data. We utilised 280 different model architectures that vary by feature selection method (S_523, f_score, fisher_score, reliefF), number of input features (200, 300, 500, 978) and algorithm (ElasticNet, Linear regression, KNN, RF, ExtraTrees, SVR, MLP). As an input for a models we utilised robust z-scores for landmark genes from L1000 dataset and as an output the ratio "Log-Ratio/Confidence interval" for BioMAP cytotox and diversity data. Robust z-scores were averaged for perturbation replicates. Input L1000 samples and output BioMAP samples were matched by drug and concentration. Also for L1000 dataset we experimented with addition of provided molecular properties (# of aromatic rings, clogp, molecular weight, TPSA) as extra features.')
        p = document.add_paragraph('Since the number of common with L1000 drugs profiled in cytotox and diversity BioMAP sets is different we used two separate output sets: cytotox and diversity. We also used 4 different training set designs to stratify the L1000 set: ')
        p = document.add_paragraph('1. Data set separated by time point (6 and 24 hours) with robust z-scores for landmark genes (Landmarks_by_time). \n\
        2. Data set separated by time point (6 and 24 hours) with robust z-scores for landmark genes and values for 4 additional molecular properties (Landmarks_and_molprop_by_time).\n\
        3. Aggregated data set with robust z-scores for landmark genes (Landmarks).\n\
        4. Aggregated data set with robust z-scores for landmark genes and values for 4 additional molecular properties (Landmarks_and_molprop).')
       
        p = document.add_paragraph('For each of the training sets feature selection methods were applied to L1000 landmark genes followed by building ML and DL models to predict the BioMAP feature. Each BioMAP biomarker was inferred independently. For each ML and DL model optimization of its hyperparameters was applied using grid search. Each model underwent 10-fold cross-validation and the resulting performance was averaged across folds. ')
        p = document.add_paragraph('In this report we present the tables with performance (mean R-squared across all biomarkers in either diversity or cytotox sets) of the models for each training set design. After each performance table we also put in separate tabs the R-squared metrics for each biomarker of top model for cytotox and top model for diversity inference. Due to curse of dimensionality and relatively small number of matching samples DNNs lost to a more conventional methods like ElasticNet and tree-based methods.')
        
        head = document.add_heading("Landmarks by time", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t1/timewise_report_land.csv')        
         
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        ################################################################
        head = document.add_heading("Top model cytotox Landmarks_by_time", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t1/Landmarks_by_time_cytotoxtop_model.csv')        
         
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        ################################################################
        head = document.add_heading("Top model diversity Landmarks_by_time", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t1/Landmarks_by_time_diversitytop_model.csv')        
         
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
                
        ################################################################
        head = document.add_heading("Landmarks_and_molprop_by_time", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t1/timewise_report_land_mol.csv')        
         
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        ################################################################
        head = document.add_heading("Top model cytotox Landmarks_and_molprop_by_time", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t1/LandmarksMol_by_time_cytotoxtop_model.csv')        
         
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        ################################################################
        head = document.add_heading("Top model diversity Landmarks_and_molprop_by_time", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t1/LandmarksMol_by_time_diversitytop_model.csv')        
         
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        
        ################################################################
        head = document.add_heading("Landmarks", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t1/report_land.csv')        
         
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        ################################################################
        head = document.add_heading("Top model cytotox Landmarks", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t1/Landmarks_cytotoxtop_model.csv')        
         
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        ################################################################
        head = document.add_heading("Top model diversity Landmarks", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t1/Landmarks_diversitytop_model.csv')        
         
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        ################################################################
        head = document.add_heading("Landmarks_and_molprop", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t1/timewise_report_land_mol.csv')        
        
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        ################################################################
        head = document.add_heading("Top model cytotox Landmarks_and_molprop", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t1/LandmarksMol_cytotoxtop_model.csv')        
        
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        ################################################################
        head = document.add_heading("Top model diversity Landmarks_and_molprop", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t1/LandmarksMol_diversitytop_model.csv')        
        
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
                
        ################################################################
        head = document.add_heading("Permutation feature importance results for Cytotox data", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t5/cytotox_pfi.csv')        
        
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        
        
        
        # end of BioMAP inference
        """
        """
        L1000-BioMAP correlation
        """ 
        """
        head = document.add_heading("L1000-BioMAP correlation", 2)
        head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph()
        
        p = document.add_paragraph("For each system (3C and SAg) and time point in L1000 dataset we matched corresponding samples to BioMAP cytotox and diversity dataset samples by drug and system. Then we calculated Pearson correlation coefficients for each L1000 gene- BioMAP marker pair. Here in the table below for each BioMAP marker we selected Top-100 most correlated (by absolute value) L1000 genes and the resulting set of Top-100 lists was aggregated.")
        
        ################################################################
        head = document.add_heading("3c_6h_cytotox", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t5/corr/corr.3c.6h.cytotox_tr.csv')        
        
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
                
        ################################################################
        head = document.add_heading("TOP-100 3c_6h_diversity", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t5/corr/corr.3c.6h.diversity_sorted.csv')        
        
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
                
        ################################################################
        head = document.add_heading("3c_24h_cytotox", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t5/corr/corr.3c.24h.cytotox_tr.csv')        
        
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        ################################################################
        head = document.add_heading("TOP-100 3c_24h_diversity", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t5/corr/corr.3c.24h.diversity_sorted.csv')        
        
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        ################################################################
        head = document.add_heading("sag_6h_cytotox", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t5/corr/corr.sag.6h.cytotox_tr.csv')        
        
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        ################################################################
        head = document.add_heading("TOP-100 sag_6h_diversity", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t5/corr/corr.sag.6h.diversity_sorted.csv')        
        
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        ################################################################
        head = document.add_heading("sag_24h_cytotox", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t5/corr/corr.sag.24h.cytotox_tr.csv')        
        
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        ################################################################
        head = document.add_heading("TOP-100 sag_24h_diversity", 4)
        df = pd.read_csv(settings.MEDIA_ROOT+'/gsk_prj4_t5/corr/corr.sag.24h.diversity_sorted.csv')        
        
        df =  pd.DataFrame(df[:50])
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        """
        # End of L1000-BioMAP correlation
        
        """
        Hit qualification
        """
        head = document.add_heading("Hit qualification", 2)
        head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph()
        
        p = document.add_paragraph('Here we utilised our already established methodology to estimate the probability of failure in clinical trials (Artemov et al., 2016). On Figures 1 and 2 you can see the overall statistic for provided L1000 dataset. In "Hit qualification table" you can see the actual score of drug perturbations analysed in this report. ')
        
        document.add_picture(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t1/hit1.png', width=Inches(5.0))
        p = document.add_paragraph('Figure 1. Probability of success for drug perturbations from provided L1000 dataset.')
        document.add_picture(settings.MEDIA_ROOT+'/report-portal/gsk_prj4_t1/hit2.png', width=Inches(5.0))
        p = document.add_paragraph('Figure 2. Probability of failure for drug perturbations from provided L1000 dataset.')
        head = document.add_heading("Hit qualification", 4)
        
        file_name = 'clinical_trial/predCT.csv'
        df = pd.read_csv(settings.MEDIA_ROOT+'/report-portal/'+report.slug+'/'+file_name)
        
        ############################## DOWN
        t = document.add_table(df.shape[0]+1, df.shape[1])            
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
                
        # End of HIT        
        
        document.save(settings.MEDIA_ROOT+'/report-pdf/'+report.slug+'.docx')
        raise Exception('stop')
        
        return context
    